"""
fix_fifth_review.py
Points 1 and 2 from fifth ChatGPT review (text edits only):

1. Section 6.2: Add one-sentence exchangeability formalization after the
   "exchangeable across the two entries" claim.
2. Section 9.5: Add "pure null" qualifier to both enrichment comparisons
   ("5.0% expected under the null" → "5.0% expected under a pure null with
   no competitor-specific differentials"; same for 0.10% figure).
"""

import shutil
from pathlib import Path
from docx import Document

DOCX_PATH = Path("judge_bias_isu_judging_system.docx")
BACKUP_PATH = Path("judge_bias_isu_judging_system.docx.bak_fifth_review")


def backup():
    shutil.copy2(DOCX_PATH, BACKUP_PATH)
    print(f"Backup → {BACKUP_PATH}")


def replace_in_para(para, old, new, label=""):
    full = para.text
    if old not in full:
        raise ValueError(f"[{label}] text not found:\n  '{old[:100]}'")
    new_text = full.replace(old, new)
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def find_para(doc, substring):
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    raise ValueError(f"Paragraph containing '{substring[:60]}' not found")


# ──────────────────────────────────────────────────────────────
# FIX 1: Section 6.2 — add exchangeability formalization sentence
# ──────────────────────────────────────────────────────────────

def fix_exchangeability_formalization(doc):
    p = find_para(doc, "leaving only judge j\u2019s idiosyncratic deviation.")
    replace_in_para(
        p,
        "leaving only judge j\u2019s idiosyncratic deviation. "
        "For each of M\u202f=\u202f10,000 permutations",
        "leaving only judge j\u2019s idiosyncratic deviation. "
        "Formally, the exchangeability assumption is that the pooled \u0394 multiset "
        "{d_A \u222a d_B} is invariant to relabeling rows as belonging to A or B "
        "under the null. "
        "For each of M\u202f=\u202f10,000 permutations",
        "fix1 exchangeability formalization"
    )
    print("  Section 6.2: added one-line exchangeability formalization.")


# ──────────────────────────────────────────────────────────────
# FIX 2: Section 9.5 — add "pure null" qualifier to enrichment phrases
# ──────────────────────────────────────────────────────────────

def fix_sec95_pure_null(doc):
    p = find_para(doc, "5.0% expected under the null; 5.07")
    # Fix both in one pass
    replace_in_para(
        p,
        "5.0% expected under the null; 5.07\u00d7\u2009enrichment), "
        "and 3.78% yield p\u2264\u200a0.001 "
        "(versus 0.10% expected; 37.8\u00d7\u2009enrichment).",
        "5.0% expected under a pure null with no competitor-specific differentials; "
        "5.07\u00d7\u2009enrichment), "
        "and 3.78% yield p\u2264\u200a0.001 "
        "(versus 0.10% under a pure null; 37.8\u00d7\u2009enrichment).",
        "fix2 sec95 pure null qualifier"
    )
    print("  Section 9.5: added 'pure null' qualifier to both enrichment comparisons.")


def main():
    backup()
    doc = Document(str(DOCX_PATH))

    print("\n--- Fix 1: Section 6.2 exchangeability formalization ---")
    fix_exchangeability_formalization(doc)

    print("\n--- Fix 2: Section 9.5 pure null qualifier ---")
    fix_sec95_pure_null(doc)

    doc.save(str(DOCX_PATH))
    print(f"\nSaved → {DOCX_PATH}")

    doc2 = Document(str(DOCX_PATH))
    all_text = " ".join(p.text for p in doc2.paragraphs)
    checks_present = [
        "pooled \u0394 multiset",
        "invariant to relabeling rows",
        "pure null with no competitor-specific differentials",
        "0.10% under a pure null",
    ]
    checks_absent = [
        "5.0% expected under the null; 5.07",
        "0.10% expected; 37.8",
    ]
    print("\nVerification:")
    all_ok = True
    for text in checks_present:
        found = text in all_text
        status = "OK" if found else "FAIL"
        if not found:
            all_ok = False
        print(f"  [{status}] present: '{text[:70]}'")
    for text in checks_absent:
        found = text in all_text
        status = "OK" if not found else "FAIL"
        if found:
            all_ok = False
        print(f"  [{status}] absent:  '{text[:70]}'")
    print(f"\n{'All checks passed.' if all_ok else 'SOME CHECKS FAILED.'}")


if __name__ == "__main__":
    main()
