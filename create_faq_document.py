#!/usr/bin/env python3
"""
create_faq_document.py

Generates OWG2026_IceDance_FD_FAQ.docx — a journalist/researcher FAQ for the
OWG 2026 Ice Dance Free Dance judging bias analysis.

Usage:
    python3 create_faq_document.py

HARDCODED VALUES — verify against figure_skating_ijs_v4.sqlite after any rerun:
    bias=+1.19, p=0.0003, q=0.034, margin=0.97, 9463 significant pairs,
    7 outcome-determinative events, 10001 permutations (M=10,000 + observed).
    Last verified: 2026-02-24 (10k rerun, seed=20260223, global CDF).
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/allman/Library/CloudStorage/Dropbox/Dropbox Mike/Judging Bias"
OUT_FILE = os.path.join(BASE, "OWG2026_IceDance_FD_FAQ.docx")

# ── Colour constants ────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x4E, 0x79)
MID_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
DARK_GREY  = RGBColor(0x40, 0x40, 0x40)
BLACK      = RGBColor(0x00, 0x00, 0x00)
RED        = RGBColor(0xC0, 0x00, 0x00)


# ── Helpers ─────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set background colour of a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    p.paragraph_format.space_after = Pt(4)


def add_subtitle(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = MID_BLUE
    p.paragraph_format.space_after = Pt(12)


def add_section_heading(doc: Document, number: int, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  {number}.  {title}  ")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Shade the paragraph (simulate a section header bar)
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "1F4E79")
    pPr.append(shading)


def add_qa(doc: Document, q: str, a: str, highlight: bool = False):
    """Add a Q + A pair. highlight=True marks the answer in dark red."""
    # Question
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before = Pt(6)
    pq.paragraph_format.space_after = Pt(2)
    pq.paragraph_format.left_indent = Inches(0.25)
    rq = pq.add_run(f"Q:  {q}")
    rq.font.bold = True
    rq.font.size = Pt(11)
    rq.font.color.rgb = DARK_BLUE

    # Answer
    pa = doc.add_paragraph()
    pa.paragraph_format.space_before = Pt(0)
    pa.paragraph_format.space_after = Pt(6)
    pa.paragraph_format.left_indent = Inches(0.5)
    ra = pa.add_run(f"A:  {a}")
    ra.font.size = Pt(10.5)
    ra.font.color.rgb = RED if highlight else DARK_GREY


def add_divider(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("─" * 90)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xBF, 0xBF, 0xBF)


# ── FAQ content ─────────────────────────────────────────────────────────────────

SECTIONS = [
    {
        "title": "About the Competition",
        "qas": [
            ("What event is this analysis about?",
             "The OWG 2026 Milano-Cortina Winter Olympics Ice Dance Free Dance, held in February 2026. "
             "Ice Dance is a discipline of figure skating performed by couples; the Free Dance is the "
             "decisive final segment."),

            ("Who competed and who won?",
             "Twenty couples competed from 14 countries. France (Laurence FOURNIER BEAUDRY / "
             "Guillaume CIZERON) won gold; USA (Madison CHOCK / Evan BATES) won silver; "
             "Canada (Piper GILLES / Paul POIRIER) won bronze."),

            ("What was the winning margin?",
             "France scored 135.64; USA scored 134.67. The winning margin was 0.97 points — "
             "less than one point."),

            ("Why does the margin matter?",
             "The statistical analysis identifies a single judge (J1 — Jezabel DABOUIS) whose "
             "marks produced a +1.19 point advantage for France over USA. Since 1.19 > 0.97, "
             "J1's scoring alone — if neutralised — would reverse the gold medal outcome."),
        ],
    },
    {
        "title": "About the ISU Scoring System",
        "qas": [
            ("How are ice dance competitors scored?",
             "Nine judges each rate every element (Grade of Execution, GOE) and three Program "
             "Component Scores (PCS): Skating Skills, Composition, and Presentation. The ISU applies "
             "a trimmed mean: the single highest and single lowest mark from the 9-judge panel are "
             "discarded, and the remaining 7 are averaged. The result is multiplied by a scale factor "
             "and summed to produce TES (technical) + PCS = TSS (total segment score)."),

            ("What is the trimmed mean and why does it exist?",
             "The ISU's trimmed mean drops the most extreme high and low mark from each panel "
             "before averaging. This is intended to reduce the influence of any one judge on "
             "the final score. However, a sufficiently biased judge can still shift the outcome "
             "if their mark is not extreme enough to be trimmed — which is exactly what the "
             "ISU-impact analysis measures."),

            ("How many judges score each event?",
             "Nine judges per event, drawn from different countries. For OWG 2026 Ice Dance Free Dance, "
             "the panel included judges from France, Russia, Spain, Finland, Germany, USA, Italy, "
             "China, and Austria."),

            ("Can judges score their own country's athletes?",
             "ISU rules do not automatically exclude judges from scoring competitors of their own "
             "nationality — they score all competitors. This creates a structural conflict of "
             "interest that has been debated in figure skating for decades. J1 in this analysis "
             "is a French judge; France won gold."),

            ("What are GOE and PCS?",
             "GOE (Grade of Execution) is an integer from -5 to +5 that each judge assigns to "
             "each skating element, reflecting execution quality. PCS (Program Component Score) "
             "rates overall skating quality across three components on a 0.25-10.00 scale. "
             "Both contribute to the total score."),
        ],
    },
    {
        "title": "About This Analysis",
        "qas": [
            ("What is the ISU-impact method?",
             'The ISU-impact method measures each judge\'s effect on the published score by asking: '
             '"What would the official total score be if this judge\'s marks were replaced by the '
             'median of the other 8 judges?" The difference — positive or negative — is the '
             "judge's impact on that competitor's score. This uses the ISU's own scoring formula "
             "exactly."),

            ("What does BiasPoints mean?",
             "BiasPoints(judge, Team A, Team B) = Impact(judge, A) − Impact(judge, B). "
             "A positive value means the judge gave Team A an ISU-score advantage over Team B "
             "through the trimmed-mean mechanism. A value of +1.19 means the judge's marks "
             "raised Team A's published score by 1.19 points relative to Team B, compared to "
             "what the panel median would have produced."),

            ("Who conducted this analysis?",
             "Michael Allman (MBA, University of Chicago Booth School of Business), with statistical "
             "AI assistance from Claude (Anthropic). The analysis uses the ISU's official "
             "'Judges Details per Skater' document as its data source."),

            ("How were results verified?",
             "The computed TES and total scores reproduce the ISU's published official scores "
             "to within 0.01 points for all 20 competitors, confirming the scoring model is exact. "
             "The full code and data are available from the author for independent verification."),

            ("What data was used?",
             "The ISU official 'Judges Details per Skater' document for OWG 2026 Milano-Cortina "
             "Ice Dance Free Dance. This is publicly released by the ISU after each competition "
             "and contains every judge's mark for every element and component."),
        ],
    },
    {
        "title": "About the Key Findings",
        "qas": [
            ("What is the headline finding?",
             "Judge J1 (Jezabel DABOUIS, France) produced a statistically significant bias of "
             "+1.19 ISU points in favour of France over USA in the gold-vs-silver pair "
             "(p = 0.0003, BH-corrected q = 0.034). France won by 0.97 points. "
             "Removing J1's effect reverses the outcome: USA wins gold.",
             True),

            ("What does 'outcome-determinative' mean?",
             "It means the identified bias is large enough to change the final standings. "
             "Specifically: J1's +1.19 point advantage to France over USA exceeds France's "
             "winning margin of 0.97 points. Under the ISU's trimmed-mean formula, "
             "recalculating without J1's contribution gives the gold medal to USA."),

            ("Who wins the gold medal if J1 is removed?",
             "USA (Madison CHOCK / Evan BATES) would win gold. J1 gave France (FOURNIER BEAUDRY / "
             "CIZERON) +0.26 points above the panel baseline and simultaneously gave USA "
             "-0.93 points below baseline — a combined differential of 1.19 points. "
             "France's winning margin was only 0.97 points.",
             True),

            ("How certain is the finding statistically?",
             "The p-value of 0.0003 means this pattern would arise by chance approximately 3 times "
             "in 10,000 simulations of a fair judge. After applying Benjamini-Hochberg false "
             "discovery rate correction across all 1,710 simultaneous tests, the q-value "
             "is 0.034 — well below the 0.05 significance threshold. This is a strong signal."),

            ("What is a p-value of 0.0003?",
             "We ran 10,000 simulations of what a fair judge's marks might look like, using "
             "a model that preserves each judge's personal scoring style. Using add-one "
             "smoothing — p = (1 + exceedances) / (10,001) — only 2 of the 10,000 "
             "permutations exceeded the observed statistic, giving "
             "p = (1 + 2) / 10,001 ≈ 0.0003 (about 3 in 10,000). That means there is "
             "roughly a 3-in-10,000 probability of seeing this pattern by chance."),

            ("Were other judges also flagged?",
             "Yes. Five of the nine judges showed at least one statistically significant pairwise "
             "differential at q ≤ 0.05 (J1, J4, J5, J6, J8). This is not surprising given "
             "that 1,710 comparisons are tested simultaneously. But only one judge's anomaly "
             "(J1) was large enough to change the gold-silver outcome. The two-part test — "
             "significant AND outcome-determinative — is what makes J1's pattern uniquely "
             "consequential."),
        ],
    },
    {
        "title": "Statistical Methods",
        "qas": [
            ("What is a permutation test?",
             "A permutation test is a simulation-based significance test. We ask: if this "
             "judge were scoring fairly (no preferential targeting), how large a bias could "
             "arise just by chance? We generate 10,000 simulations of a fair judge and "
             "measure how often they produce a bias as large as what we observed. "
             "The fraction of times they do is the p-value."),

            ("What is the 'style-adjusted' null model?",
             "Judges have personal styles — some score more generously, others more strictly. "
             "Our null model preserves each judge's personal scoring distribution while "
             "removing any preferential targeting. We convert each mark to a percentile "
             "within that judge's scoring distribution, then randomly reassign which teams "
             "receive which percentile. This tests: is the judge's pattern specific to "
             "these two teams, or just their normal style?"),

            ("How much scoring data does the analysis use for each judge?",
             "The method builds each judge's scoring baseline from every mark they gave "
             "across all events they judged in the database. For the OWG 2026 ice dance "
             "panel, the nine judges each appear in 1–4 events at the 2026 Olympics, "
             "providing 184–396 GOE marks per judge. The broader database covers 142 events "
             "across 17 major ISU competitions (2022–2026), and for judges who appear at "
             "multiple competitions, the baseline draws on all of them.\n\n"
             "The permutation test asks: given THIS judge's observed scoring distribution, "
             "is what she gave France-vs-USA in this specific event an outlier? "
             "The answer for J1 is yes — strongly. This is important because a judge who "
             "tends to give generous marks to everyone is NOT flagged for giving France a "
             "high mark; that's just their style. What IS flagged is a judge giving one team "
             "marks that are systematically above her own pattern, while simultaneously "
             "scoring the rival team systematically below her pattern."),

            ("What is the Benjamini-Hochberg correction, and why does it matter here?",
             "When you run 1,710 statistical tests simultaneously (9 judges × 190 team pairs), "
             "pure chance will make roughly 85 of them look significant at the p < 0.05 level "
             "— even if every single judge is perfectly fair. This is the 'multiple comparisons "
             "problem': the more you look, the more false alarms you get.\n\n"
             "The Benjamini-Hochberg False Discovery Rate method (BH-FDR) controls this. "
             "It guarantees that among all the findings we call 'significant,' the expected "
             "proportion of false discoveries is no more than the threshold we set (5%). "
             "In other words, if we repeatedly applied this procedure, on average no more "
             "than 5% of our flagged results would be false alarms.\n\n"
             "Mechanically, BH works by ranking all 1,710 p-values from smallest to largest, "
             "then progressively raising the bar for what counts as significant. Only tests "
             "with the strongest signals survive this gauntlet. The output is a q-value for "
             "each test. J1's FRA-USA q = 0.034 — meaning even after this strict "
             "multi-test correction, the finding survives.\n\n"
             "In plain English: we tested 1,710 combinations. The BH procedure identified "
             "21 significant pairs across 5 judges. By the FDR guarantee, the expected "
             "false discovery rate among those 21 is controlled at a low level."),

            ("What is C(20,2) = 190?",
             "With 20 competitors, there are C(20,2) = 20 × 19 ÷ 2 = 190 unique pairs of "
             "teams. Each judge's marks imply a preference for every possible pairing — "
             "did they score Team A higher or lower than Team B? With 9 judges, that's "
             "9 × 190 = 1,710 simultaneous comparisons tested."),

            ("Why use 10,000 permutations?",
             "More permutations give a more precise p-value estimate. With 10,000 permutations "
             "and a smoothed formula p = (1 + exceedances) / (10,001), the smallest "
             "possible p-value is approximately 0.0001. Our key finding has p = 0.0003, "
             "well above the floor, ensuring the estimate is reliable."),

            ("What is the method version stored in the database?",
             "isuimpact_quantile_v1. RNG seed: 20260223 (numpy.random.default_rng). "
             "CDF scope: global (all available marks per judge in the database). "
             "This ensures full reproducibility."),
        ],
    },
    {
        "title": "About the Judges / Conflict of Interest",
        "qas": [
            ("Who is J1 — Jezabel DABOUIS?",
             "Jezabel DABOUIS is a French ice dance judge assigned to the OWG 2026 Ice Dance "
             "Free Dance panel. In the ISU system, judges are assigned by their national "
             "federations and can score all competitors, including those from their own country."),

            ("Is it unusual for a judge to show this pattern?",
             "Among the 9 judges on this panel, J1 is one of 5 showing significant patterns. "
             "However, J1 is unique in that their most significant pair is the gold-silver "
             "matchup — and they are a national of the gold-winning country. The combination "
             "of statistical significance, outcome-determinative magnitude, and national "
             "affiliation makes this pattern particularly noteworthy."),

            ("Does the ISU prohibit judges from scoring their own nationals?",
             "ISU rules do not automatically prohibit judges from scoring athletes from their "
             "own federation. Some critics argue this is a structural flaw; others defend it "
             "on the grounds that professional judges can be objective. The ISU does prohibit "
             "judges from having 'personal or financial interest' in the result, but "
             "national affiliation alone is not sufficient grounds for exclusion."),

            ("Is this evidence of intentional misconduct?",
             "No. Statistical analysis can detect patterns that are improbable under a "
             "fair-judging assumption — it cannot determine intent. The finding is that "
             "J1's marks produced a significantly biased outcome. Whether this reflects "
             "unconscious national preference, deliberate scoring manipulation, or some "
             "other factor cannot be determined from marks alone."),
        ],
    },
    {
        "title": "Limitations & Caveats",
        "qas": [
            ("Could this finding be a coincidence?",
             "At p = 0.0003 and q = 0.034 (after correcting for 1,710 simultaneous tests), "
             "the finding is statistically significant by standard scientific criteria. "
             "However, statistical significance is not the same as certainty. The q-value "
             "of 0.034 means that if we treat all findings with q ≤ 0.034 as significant, "
             "the expected false discovery rate among them is about 3.4%. This is a "
             "statement about the procedure's long-run accuracy, not a probability "
             "attached to any single finding."),

            ("Could there be innocent explanations for J1's pattern?",
             "Possibly. J1 may genuinely believe the French skaters deserved higher marks. "
             "Judges are human and bring subjective aesthetic preferences to their scoring. "
             "The style-adjusted null model controls for global scoring tendencies (generous "
             "vs. strict), but cannot rule out legitimate aesthetic differences in how a "
             "judge values specific technical or artistic elements."),

            ("Does this prove the gold medal should be reassigned?",
             "No. Statistical evidence of bias does not automatically trigger a result "
             "reversal under ISU rules. Olympic results are governed by the ISU Constitution "
             "and the IOC charter, which have specific procedures for protests and appeals. "
             "This analysis provides evidence that warrants investigation; it is not itself "
             "a finding of misconduct."),

            ("Is this a single-event finding?",
             "Yes. This analysis covers one event: OWG 2026 Ice Dance Free Dance. "
             "A broader database of 142 figure skating events exists and could be analysed "
             "to identify whether any judge shows consistent cross-event patterns. "
             "That analysis has not yet been published."),

            ("Are there alternative statistical methods that might give different results?",
             "Yes. Different null models, significance thresholds, or bias metrics could "
             "produce different results. The ISU-impact method was chosen because it uses "
             "the ISU's own scoring formula exactly, making results directly interpretable "
             "in competition points. Alternative approaches (e.g., rank-based tests, "
             "regression models) generally yield corroborating findings."),
        ],
    },
    {
        "title": "Next Steps & Broader Context",
        "qas": [
            ("What can the ISU do about this?",
             "The ISU could: (1) investigate the specific finding through its ethics process; "
             "(2) implement automatic exclusion of judges scoring their own nationals; "
             "(3) publish fuller statistical monitoring of judge panels after each major "
             "competition; (4) consider algorithmic flags for judges showing systematic "
             "patterns across events."),

            ("Has this type of bias been found before in figure skating?",
             "National judging bias in figure skating has been documented in academic "
             "literature since at least the 2002 Salt Lake City scandal (pairs skating). "
             "The ISU introduced the current anonymous randomised-panel system partly in "
             "response. This analysis applies a new, more rigorous method to the current "
             "system and finds that bias has not been eliminated."),

            ("Is this sufficient evidence for an official appeal?",
             "Under current ISU rules, an appeal must be filed by a competing federation "
             "within a short window after the competition. Whether the statistical evidence "
             "in this analysis would meet the evidentiary standard for a successful appeal "
             "is a legal and procedural question beyond the scope of this analysis. "
             "Counsel with expertise in sports arbitration would be needed."),

            ("What reforms would prevent this kind of outcome?",
             "Proposed reforms include: excluding judges from scoring their own nationals; "
             "real-time statistical monitoring of judge deviations from the panel median; "
             "publishing individual judge marks alongside competition results (currently "
             "available but not prominently displayed); and requiring judges to certify "
             "no conflict of interest before each event."),

            ("Where can I get the full data and analysis?",
             "The complete analysis workbook (OWG 2026 Ice Dance Free Dance — Complete "
             "Analysis Workbook) is available from the analyst. It contains 12 tabs: "
             "ISU official scoring, pairwise bias statistics for all 9 judges and all "
             "190 team pairs, and full methodology documentation."),
        ],
    },
]


# ── Main ─────────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── Title block ──
    add_title(doc, "Frequently Asked Questions")
    add_title(doc, "OWG 2026 Ice Dance Free Dance — Judging Bias Analysis")
    add_subtitle(doc,
        "For journalists, researchers, and the public  ·  "
        "Analysis by Michael Allman (MBA, University of Chicago Booth School of Business)")

    p = doc.add_paragraph()
    run = p.add_run(
        "This document addresses common questions about the statistical analysis of potential "
        "judging bias in the OWG 2026 Milano-Cortina Ice Dance Free Dance. "
        "The headline finding: Judge J1 (Jezabel DABOUIS, France) produced a statistically "
        "significant +1.19 point advantage for France over USA (p = 0.0003). "
        "France won by 0.97 points. Removing J1's effect reverses the gold medal."
    )
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK_GREY
    run.font.italic = True
    p.paragraph_format.space_after = Pt(14)

    add_divider(doc)

    # ── Sections ──
    for i, section_data in enumerate(SECTIONS, start=1):
        add_section_heading(doc, i, section_data["title"])
        for qa in section_data["qas"]:
            highlight = len(qa) > 2 and qa[2]
            add_qa(doc, qa[0], qa[1], highlight=highlight)
        add_divider(doc)

    # ── Footer note ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Analysis by Michael Allman · MBA, University of Chicago Booth School of Business · "
        "Statistical assistance: Claude (Anthropic) · February 2026\n"
        "Data source: ISU official Judges Details per Skater, OWG 2026 Milano-Cortina · "
        "Method: isuimpact_quantile_v1 · Seed: 20260223 · 10,000 permutations"
    )
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(OUT_FILE)
    print(f"Saved: {OUT_FILE}")
    n_qas = sum(len(s["qas"]) for s in SECTIONS)
    print(f"  {len(SECTIONS)} categories, {n_qas} Q&A pairs")


if __name__ == "__main__":
    main()
