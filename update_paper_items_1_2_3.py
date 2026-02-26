"""
update_paper_items_1_2_3.py
Adds three remaining paper elements:
  1. Dataset descriptor table in Section 3 (after 3.3 paragraph)
  2. Calibration results in Section 9.5 (replaces "deferred" language)
  3. J1 nationality analysis paragraph + table in Section 8 (after Table 1 note)
"""

import copy
import shutil
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX_PATH = Path("judge_bias_isu_judging_system.docx")
BACKUP_PATH = Path("judge_bias_isu_judging_system.docx.bak_items123")


def backup():
    shutil.copy2(DOCX_PATH, BACKUP_PATH)
    print(f"Backup → {BACKUP_PATH}")


def make_paragraph_after(ref_element, style_name="Normal"):
    """Create an empty paragraph element positioned right after ref_element."""
    p = OxmlElement("w:p")
    ref_element.addnext(p)
    return p


def set_para_text(p_elem, text, bold=False, italic=False, style=None):
    """Add a run with text to an OxmlElement paragraph."""
    r = OxmlElement("w:r")
    if bold or italic:
        rPr = OxmlElement("w:rPr")
        if bold:
            b = OxmlElement("w:b")
            rPr.append(b)
        if italic:
            i = OxmlElement("w:i")
            rPr.append(i)
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    p_elem.append(r)
    return p_elem


def make_tbl_row(values, header=False):
    """Return a <w:tr> element with cells for each value in values."""
    tr = OxmlElement("w:tr")
    for val in values:
        tc = OxmlElement("w:tc")
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        if header:
            rPr = OxmlElement("w:rPr")
            b = OxmlElement("w:b")
            rPr.append(b)
            r.append(rPr)
        t = OxmlElement("w:t")
        t.text = str(val)
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        p.append(r)
        tc.append(p)
        tr.append(tc)
    return tr


def build_table_element(headers, rows):
    """Build a minimal <w:tbl> XML element from headers + data rows."""
    tbl = OxmlElement("w:tbl")

    # tblPr
    tblPr = OxmlElement("w:tblPr")
    tblStyle = OxmlElement("w:tblStyle")
    tblStyle.set(qn("w:val"), "TableGrid")
    tblPr.append(tblStyle)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "0")
    tblW.set(qn("w:type"), "auto")
    tblPr.append(tblW)
    tbl.append(tblPr)

    # tblGrid (one col per header, equal widths)
    tblGrid = OxmlElement("w:tblGrid")
    n_cols = len(headers)
    for _ in range(n_cols):
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(int(9360 / n_cols)))
        tblGrid.append(gridCol)
    tbl.append(tblGrid)

    # header row
    tbl.append(make_tbl_row(headers, header=True))

    # data rows
    for row in rows:
        tbl.append(make_tbl_row(row))

    return tbl


def find_para_by_text(doc, substring):
    """Return first paragraph whose text contains substring, or None."""
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    return None


# ─────────────────────────────────────────────
# TASK 1: Dataset table in Section 3
# ─────────────────────────────────────────────

def add_dataset_table(doc):
    # Find the paragraph containing Section 3.3
    p = find_para_by_text(doc, "3.3. Multi-event database")
    assert p is not None, "Could not find '3.3. Multi-event database'"

    # Update 3.3 prose: mention 144 total and 2 excluded
    old_text = ("spanning 17 competitions and 142 discipline\u00d7segment events, "
                "containing 206,682 individual GOE grades and 84,922 individual PCS marks "
                "(291,604 individual judge marks in total). "
                "The present paper focuses on the single-event methodology and its "
                "validation on the Olympic example; aggregate results across the full "
                "database are reported in Section 9.")
    new_text = ("spanning 17 competitions across five seasons (2021/22\u20132025/26) and "
                "144 events in total — 142 analyzed events and 2 excluded events (the ISU "
                "Four Continents Championships 2022 Ice Dance Rhythm Dance and Free Dance, "
                "which used non-standard 7-judge panels). The database contains 206,682 "
                "individual GOE grades and 84,922 individual PCS marks (291,604 total judge "
                "marks) from 2,686 competitor entries. Table\u00a0B provides a complete "
                "breakdown; aggregate results are reported in Section\u00a09.")

    # Rebuild the paragraph runs with the replacement
    # The paragraph has mixed runs; easiest is to find the run containing the old text
    full_text = p.text
    if old_text not in full_text:
        print(f"WARNING: old 3.3 text not found verbatim; checking partial match...")
        print(f"  Looking for: ...{old_text[:60]}...")
        print(f"  Found para: ...{full_text[-200:]}")
        raise ValueError("3.3 prose replacement text not found")

    # The paragraph may span multiple runs; rebuild it
    new_full = full_text.replace(old_text, new_text)
    # Clear all runs
    for run in p.runs:
        run.text = ""
    # Set text in first run
    if p.runs:
        p.runs[0].text = new_full
    else:
        p.add_run(new_full)

    print("  Updated 3.3 prose.")

    # Now insert Table B title + table after p
    ref_elem = p._element

    # Table caption paragraph
    cap_p = OxmlElement("w:p")
    ref_elem.addnext(cap_p)
    set_para_text(cap_p, "Table B. Database characteristics.")

    # Dataset table
    dataset_headers = ["Characteristic", "Count / Value"]
    dataset_rows = [
        ["Seasons", "2021/22 \u2013 2025/26"],
        ["ISU competitions", "17"],
        ["Events (total)", "144"],
        ["\u2003Main discipline events analyzed", "134"],
        ["\u2003Team event segments analyzed", "8"],
        ["\u2003Excluded (7-judge panels)", "2"],
        ["Competitor entries", "2,686"],
        ["Panel judges (distinct roles)", "1,278"],
        ["Individual judge marks", "291,604"],
        ["Pairwise permutation tests (v2)", "271,728"],
    ]
    tbl = build_table_element(dataset_headers, dataset_rows)
    cap_p.addnext(tbl)

    # Add blank paragraph after table (standard docx practice)
    blank_p = OxmlElement("w:p")
    tbl.addnext(blank_p)

    print("  Inserted Table B (dataset characteristics).")


# ─────────────────────────────────────────────
# TASK 2: Replace calibration placeholder in 9.5
# ─────────────────────────────────────────────

def update_calibration(doc):
    p = find_para_by_text(doc, "9.5. Stability, calibration, and power")
    assert p is not None, "Could not find Section 9.5"

    old_calib = ("Calibration: the 19.0% of events with zero significant pairs is a "
                 "coarse consistency check; a formal p-value histogram across null-like "
                 "events is deferred; code to generate it is available in the project's "
                 "GitHub repository.")
    new_calib = ("Calibration: residual-label permutation p-values are exact by "
                 "construction \u2014 under the null, the permutation p-value is "
                 "discrete-uniform on {1/M, 2/M, \u2026, 1} (converging to continuous "
                 "uniform(0,\u202f1) as M\u2192\u221e), so deviations from uniformity reflect "
                 "genuine signal rather than test miscalibration. Across 271,728 pairwise "
                 "tests, 25.4% yield p\u2264\u200a0.05 (versus 5.0% expected under the null; "
                 "5.07\u00d7\u2009enrichment), and 3.78% yield p\u2264\u200a0.001 (versus 0.10% "
                 "expected; 37.8\u00d7\u2009enrichment). The p-value distribution shows a "
                 "pronounced spike at small values and approximately uniform density above "
                 "p\u2248\u200a0.05, the two-component signature of a mixture of true-null and "
                 "true-alternative hypotheses. The 7.4% BH-FDR discovery rate (q\u2264\u200a0.05) "
                 "is conservative relative to the nominal 5.0% by design, consistent with a "
                 "non-trivial fraction of true alternatives.")

    old_header = ("9.5. Stability, calibration, and power (prospective). Three validation "
                  "analyses are deferred to future work.")
    new_header = ("9.5. Stability, calibration, and power (prospective). Two of three "
                  "validation analyses are deferred to future work.")

    full_text = p.text
    # Apply both replacements
    if old_calib not in full_text:
        print(f"WARNING: calibration placeholder not found verbatim")
        # Print what we have
        idx = full_text.find("Calibration:")
        if idx >= 0:
            print(f"  Found 'Calibration:' at position {idx}: ...{full_text[idx:idx+200]}...")
        raise ValueError("Calibration text not found for replacement")

    if old_header not in full_text:
        print(f"WARNING: 9.5 header 'Three validation' not found verbatim")
        raise ValueError("9.5 header not found for replacement")

    new_text = full_text.replace(old_header, new_header).replace(old_calib, new_calib)

    # Rebuild runs
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)

    print("  Updated Section 9.5 calibration text.")


# ─────────────────────────────────────────────
# TASK 3: J1 nationality analysis in Section 8
# ─────────────────────────────────────────────

def add_j1_nationality(doc):
    # Insert after the Table 1 note paragraph (para [30])
    # That paragraph starts with "Note: Positive BiasPoints indicate..."
    p = find_para_by_text(doc, "Note: Positive BiasPoints indicate a margin shift toward the 1st-place team")
    assert p is not None, "Could not find Table 1 note paragraph"

    ref_elem = p._element

    # Nationality pattern paragraph
    nat_text = (
        "J1 nationality pattern. Table\u00a0C presents J1\u2019s estimated impact on "
        "every team\u2019s total segment score in this event. J1 (Jezabel Dabouis, FRA) "
        "provided a net positive adjustment to both French-represented teams "
        "(+0.26\u2009pts for gold-medal team Fournier Beaudry / Cizeron; "
        "+0.08\u2009pts for Lopareva / Brissaud), while all three USA-represented teams "
        "received negative adjustments (\u22120.93\u2009pts for silver-medal Chock / Bates; "
        "\u22120.43\u2009pts for Zingas / Kolesnik; \u22120.34\u2009pts for Carreira / "
        "Ponomarenko). The directional asymmetry between FRA and USA entries is "
        "consistent with the significant BiasPoints detected for the gold\u2013silver pair "
        "(B_{J1}(FRA,\u2009USA)\u2009=\u2009+1.19, q\u2009=\u20090.012). "
        "No other nationality grouping shows a comparable unidirectional pattern in "
        "J1\u2019s impact column."
    )

    nat_p = OxmlElement("w:p")
    ref_elem.addnext(nat_p)
    set_para_text(nat_p, nat_text)

    # Table C caption
    cap_p = OxmlElement("w:p")
    nat_p.addnext(cap_p)
    set_para_text(cap_p, "Table C. J1 estimated impact on each team\u2019s TSS (OWG 2026 Ice Dance Free Dance).")

    # J1 impact data from DB query (all 20 teams, sorted by final rank)
    j1_data = [
        # rank, team, noc, impact
        (1,  "Fournier Beaudry / Cizeron",     "FRA", "+0.26"),
        (2,  "Chock / Bates",                   "USA", "\u22120.93"),
        (3,  "Gilles / Poirier",                "CAN", "\u22120.68"),
        (4,  "Guignard / Fabbri",               "ITA", "\u22120.39"),
        (5,  "Zingas / Kolesnik",               "USA", "\u22120.43"),
        (6,  "Smart / Dieck",                   "ESP", "\u22120.35"),
        (7,  "Reed / Ambrulevicius",            "LTU", "+0.33"),
        (8,  "Lopareva / Brissaud",             "FRA", "+0.08"),
        (9,  "Lajoie / Lagha",                  "CAN", "+0.24"),
        (10, "Carreira / Ponomarenko",          "USA", "\u22120.34"),
        (11, "Davis / Smolkin",                 "GEO", "+0.14"),
        (12, "Fear / Gibson",                   "GBR", "\u22120.05"),
        (13, "Turkkila / Versluis",             "FIN", "+0.05"),
        (14, "Lauriault / le Gac",              "CAN", "+0.29"),
        (15, "Taschlerova / Taschler",          "CZE", "+0.24"),
        (16, "Mrazkova / Mrazek",               "CZE", "\u22120.30"),
        (17, "Harris / Chan",                   "AUS", "\u22120.21"),
        (18, "Bekker / Hernandez",              "GBR", "+0.04"),
        (19, "Val / Kazimov",                   "ESP", "+0.43"),
        (20, "Reitan / Majorov",                "SWE", "+0.46"),
    ]

    tbl_headers = ["Final rank", "Team", "NOC", "J1 impact (pts)"]
    tbl_rows = [
        [str(rank), team, noc, impact]
        for rank, team, noc, impact in j1_data
    ]
    tbl = build_table_element(tbl_headers, tbl_rows)
    cap_p.addnext(tbl)

    # Blank paragraph after table
    blank_p = OxmlElement("w:p")
    tbl.addnext(blank_p)

    # Summary note paragraph
    note_text = (
        "Note: Impact points are estimated via median-of-others neutralization "
        "(isuimpact_residual_v1, M\u2009=\u200910,000, seed\u2009=\u200920260223). "
        "FRA entries: mean +0.17\u2009pts. USA entries: mean \u22120.57\u2009pts."
    )
    note_p = OxmlElement("w:p")
    blank_p.addnext(note_p)
    set_para_text(note_p, note_text)

    print("  Inserted J1 nationality analysis paragraph and Table C.")


def main():
    backup()
    doc = Document(str(DOCX_PATH))

    print("\n--- Task 1: Dataset table (Section 3) ---")
    add_dataset_table(doc)

    print("\n--- Task 2: Calibration results (Section 9.5) ---")
    update_calibration(doc)

    print("\n--- Task 3: J1 nationality analysis (Section 8) ---")
    add_j1_nationality(doc)

    doc.save(str(DOCX_PATH))
    print(f"\nSaved → {DOCX_PATH}")

    # Quick verification
    doc2 = Document(str(DOCX_PATH))
    print(f"\nVerification:")
    print(f"  Total paragraphs: {len(doc2.paragraphs)}")
    print(f"  Total tables: {len(doc2.tables)}")

    checks = [
        "Table B. Database characteristics",
        "271,728 pairwise permutation tests",
        "25.4% yield p",
        "Two of three validation analyses",
        "J1 nationality pattern",
        "Table C. J1 estimated impact",
        "mean +0.17",
    ]
    for c in checks:
        found = any(c in p.text for p in doc2.paragraphs)
        status = "OK" if found else "MISSING"
        print(f"  [{status}] '{c[:60]}'")


if __name__ == "__main__":
    main()
