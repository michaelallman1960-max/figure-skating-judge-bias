"""
fix_seventh_review.py
Seventh-review comprehensive fixes:
A1, A2, A3(§9.5), A4, A5, A6(ISU names + American English), A7(exact p-values),
A8, B1, B2(CF margin column), B3, C1/C5/C6(future-work para), C4, D2, E2
"""

import shutil
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX = Path("judge_bias_isu_judging_system.docx")
BAK  = Path("judge_bias_isu_judging_system.docx.bak_seventh_review")
W    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# ── helpers ─────────────────────────────────────────────────────────────────

def replace_in_para(para, old, new):
    full = para.text
    if old not in full:
        raise ValueError(f"NOT FOUND: {repr(old[:100])}")
    new_text = full.replace(old, new, 1)
    for i, run in enumerate(para.runs):
        run.text = new_text if i == 0 else ""


def insert_para_after(after_para, text):
    new_p = deepcopy(after_para._element)
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    for ins in new_p.findall(qn("w:ins")):
        new_p.remove(ins)
    r_elem = OxmlElement("w:r")
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_elem.append(t_elem)
    new_p.append(r_elem)
    after_para._element.addnext(new_p)


def set_cell_text(tc, text):
    for t_elem in tc.findall(f".//{{{W}}}t"):
        t_elem.text = ""
    t_elems = tc.findall(f".//{{{W}}}t")
    if t_elems:
        t_elems[0].text = text


def add_table_column(table, header_text, cell_texts):
    tbl = table._tbl
    rows = [ch for ch in tbl if ch.tag == f"{{{W}}}tr"]
    tblGrid = tbl.find(f"{{{W}}}tblGrid")
    if tblGrid is not None:
        last_gc = tblGrid.findall(f"{{{W}}}gridCol")[-1]
        tblGrid.append(deepcopy(last_gc))
    all_texts = [header_text] + list(cell_texts)
    for ri, row in enumerate(rows):
        cells = [ch for ch in row if ch.tag == f"{{{W}}}tc"]
        new_tc = deepcopy(cells[-1])
        for t_elem in new_tc.findall(f".//{{{W}}}t"):
            t_elem.text = ""
        t_elems = new_tc.findall(f".//{{{W}}}t")
        if t_elems:
            t_elems[0].text = all_texts[ri] if ri < len(all_texts) else ""
        row.append(new_tc)


# ── main ────────────────────────────────────────────────────────────────────

def main():
    shutil.copy(DOCX, BAK)
    print(f"Backed up to {BAK}")
    doc = Document(DOCX)

    # ── A1: §6.2 "removes" → "substantially attenuates" ─────────────────────
    for p in doc.paragraphs:
        if "neutralization removes the shared competitor-quality signal" in p.text:
            replace_in_para(p,
                "neutralization removes the shared competitor-quality signal row by row, "
                "leaving only judge j\u2019s idiosyncratic deviation.",
                "neutralization substantially attenuates the shared competitor-quality signal "
                "row by row, leaving primarily judge j\u2019s idiosyncratic deviation.")
            print("  \u2713 A1: 'removes' \u2192 'substantially attenuates'")
            break
    else:
        raise ValueError("A1: text not found")

    # ── A2: §6.2 validity sentence ───────────────────────────────────────────
    for p in doc.paragraphs:
        if "within-entry dependence affects power and the discreteness" in p.text:
            replace_in_para(p,
                "the residual-label test requires only exchangeability of the pooled "
                "\u0394 multiset across entries under the null; within-entry dependence "
                "affects power and the discreteness of the permutation distribution, "
                "not test validity.",
                "the residual-label test is exact conditional on the exchangeability "
                "of the pooled \u0394 multiset under relabeling. Within-entry dependence "
                "is permissible provided the dependence structure is symmetric across "
                "A and B; if dependence differs systematically between A and B (e.g., "
                "due to differing row-type mixes or element leverage), exchangeability "
                "may be violated.")
            print("  \u2713 A2: validity sentence updated")
            break
    else:
        raise ValueError("A2: text not found")

    # ── A3 §9.5: soften "genuine signal" ─────────────────────────────────────
    for p in doc.paragraphs:
        if "deviations from uniformity reflect genuine signal rather than test miscalibration" in p.text:
            replace_in_para(p,
                "so deviations from uniformity reflect genuine signal rather than "
                "test miscalibration.",
                "so deviations from uniformity may reflect a mixture of true alternatives "
                "and residual departures from exchangeability; calibration checks are "
                "therefore important.")
            print("  \u2713 A3: \u00a79.5 calibration claim softened")
            break
    else:
        raise ValueError("A3: text not found")

    # ── A4: §9.1 "indicating" → "consistent with a mixture" ─────────────────
    for p in doc.paragraphs:
        if "indicating systematic detection of real signal in the data" in p.text:
            replace_in_para(p,
                "\u2014 indicating systematic detection of real signal in the data.",
                "\u2014 consistent with a mixture of true alternatives and null tests.")
            print("  \u2713 A4: \u00a79.1 enrichment language softened")
            break
    else:
        raise ValueError("A4: text not found")

    # ── A5: §9.5 "anonymization" → identity normalization ───────────────────
    for p in doc.paragraphs:
        if "ISU judge anonymization complicates this for most events" in p.text:
            replace_in_para(p,
                "ISU judge anonymization complicates this for most events in the "
                "current database, though the named OWG 2026 panel is an exception.",
                "ISU judging sheets include judge names, but cross-event identity "
                "linking requires name normalization across events; stability "
                "analysis is deferred to future work.")
            print("  \u2713 A5: 'anonymization' \u2192 identity normalization")
            break
    else:
        raise ValueError("A5: text not found")

    # ── A6a: ISU discipline names in §9.2 text ───────────────────────────────
    for p in doc.paragraphs:
        if "Women's singles showed the highest proportion" in p.text:
            replace_in_para(p,
                "Women's singles showed the highest proportion (9.0% of tests), "
                "followed by Men's singles (7.2%), Pair skating (7.0%), and Ice Dance (6.0%).",
                "Women Single Skating showed the highest proportion (9.0% of tests), "
                "followed by Men Single Skating (7.2%), Pair Skating (7.0%), and Ice Dance (6.0%).")
            print("  \u2713 A6a: \u00a79.2 discipline names \u2192 ISU official")
            break
    else:
        raise ValueError("A6a: text not found")

    # ── A6b: "programme" → "Program Component Score" in §8 ──────────────────
    for p in doc.paragraphs:
        if "programme component scores (PCS)" in p.text:
            replace_in_para(p,
                "programme component scores (PCS)",
                "Program Component Score (PCS)")
            print("  \u2713 A6b: 'programme' \u2192 'Program Component Score'")
            break
    else:
        raise ValueError("A6b: text not found")

    # ── A8: soften "No other nationality grouping" ───────────────────────────
    for p in doc.paragraphs:
        if "No other nationality grouping shows a comparable unidirectional pattern" in p.text:
            replace_in_para(p,
                "No other nationality grouping shows a comparable unidirectional "
                "pattern in J1\u2019s impact column.",
                "We did not observe similarly unidirectional patterns among other "
                "NOC groupings in this segment.")
            print("  \u2713 A8: J1 nationality claim softened")
            break
    else:
        raise ValueError("A8: text not found")

    # ── B1: LOJO caution sentence ─────────────────────────────────────────────
    for p in doc.paragraphs:
        if "462 podium changes (35.9%)." in p.text:
            replace_in_para(p,
                "462 podium changes (35.9%).",
                "462 podium changes (35.9%). LOJO is a sensitivity analysis, not an "
                "inference procedure; high flip rates can occur in close competitions "
                "even under benign stylistic dispersion, and should not be interpreted "
                "as evidence of misconduct.")
            print("  \u2713 B1: LOJO caution sentence added")
            break
    else:
        raise ValueError("B1: text not found")

    # ── B3: LOJO/ISU-impact discordance explanation ──────────────────────────
    DISCORDANCE = (
        " The nine outcome-determinative events (Table\u202f3) include five where LOJO "
        "independently confirms a change in gold-medal winner (LOJO\u202f=\u202fYes) and four "
        "where LOJO does not flip the gold despite the pairwise bias exceeding the "
        "margin (LOJO\u202f=\u202fNo). The discordance reflects a structural difference: "
        "ISU-impact measures judge\u202fj\u2019s differential on the specific gold\u2013silver "
        "pair under median-of-others neutralization, while LOJO replaces j\u2019s marks "
        "for all competitors and recomputes the full trimmed-mean standings. A judge "
        "whose bias is concentrated on the gold\u2013silver pair may have countervailing "
        "effects on other entries \u2014 or may shift the panel\u2019s trimming set \u2014 such that "
        "the full LOJO recomputation does not flip the winner even when the pairwise "
        "criterion is satisfied. The counterfactual margins in Table\u202f3 "
        "(LOJO\u202fCF\u202f\u0394) show the gold\u2013silver gap under the LOJO scenario; for "
        "LOJO\u202f=\u202fNo rows, some values differ markedly from the actual margin "
        "(e.g., Europeans\u202f2022 Men Single Skating: 0.09\u202fpts counterfactual vs. "
        "0.70\u202fpts actual), illustrating near-determinative influence even when the "
        "gold-medal winner does not formally change."
    )
    for p in doc.paragraphs:
        if "convergent evidence across two methodologically distinct frameworks." in p.text:
            replace_in_para(p,
                "convergent evidence across two methodologically distinct frameworks.",
                "convergent evidence across two methodologically distinct frameworks." + DISCORDANCE)
            print("  \u2713 B3: LOJO/ISU-impact discordance explanation added")
            break
    else:
        raise ValueError("B3: text not found")

    # ── C4: 24-event calibration note in §9.5 ───────────────────────────────
    for p in doc.paragraphs:
        if "consistent with a non-trivial fraction of true alternatives." in p.text:
            replace_in_para(p,
                "consistent with a non-trivial fraction of true alternatives.",
                "consistent with a non-trivial fraction of true alternatives. "
                "The 24 events (16.9%) with zero BH-significant pairs (Section\u202f9.1) "
                "form a natural near-null calibration subset; a p-value histogram or "
                "QQ-plot restricted to those events is a planned robustness check to "
                "confirm that the permutation test is well-calibrated under near-null conditions.")
            print("  \u2713 C4: 24-event calibration check added to \u00a79.5")
            break
    else:
        raise ValueError("C4: text not found")

    # ── C1/C5/C6: future-work paragraph after §10 Discussion ─────────────────
    FUTURE_WORK = (
        "Planned robustness checks and extensions include: "
        "(i)\u202fStratified residual-label permutation \u2014 repeating the analysis with "
        "GOE rows and PCS rows as separate within-category tests \u2014 would provide "
        "formal p-values for each score component and directly test whether pooling "
        "the two row types drives results. For the OWG\u202f2026 flagship result, J1\u2019s "
        "+1.19\u202fpt net margin shift decomposes as approximately +1.01\u202fpts (GOE) and "
        "+0.20\u202fpts (PCS; Section\u202f8); stratified permutation would confirm or qualify "
        "whether each component independently reaches significance. "
        "(ii)\u202fA database-wide hierarchical FDR (Benjamini\u2013Yekutieli) sensitivity "
        "analysis across all 271,728 tests is planned as a supplement to the "
        "event-local BH procedure reported here. "
        "(iii)\u202fMulti-judge collusion detection \u2014 computing cross-judge correlations "
        "of impact vectors I_j(T) within events, clustering correlated judges, and "
        "testing whether judge clusters collectively shift margins \u2014 addresses forms "
        "of bias that single-judge tests are not optimized to detect."
    )
    disc_para = None
    for p in doc.paragraphs:
        if "Finally, the methodology is event-local." in p.text:
            disc_para = p
            break
    if disc_para is None:
        raise ValueError("C1/C5/C6: Discussion anchor not found")
    insert_para_after(disc_para, FUTURE_WORK)
    print("  \u2713 C1/C5/C6: Future-work paragraph added after \u00a710 Discussion")

    # ── E2: "remaining eight judges" → "remaining J−1 judges" in §9.4 ───────
    for p in doc.paragraphs:
        if "remaining eight judges' marks" in p.text:
            replace_in_para(p,
                "remaining eight judges' marks",
                "remaining J\u22121 judges' marks")
            print("  \u2713 E2: 'remaining eight' \u2192 'remaining J\u22121' in \u00a79.4")
            break
    else:
        raise ValueError("E2: text not found")

    # ── E2b: add J=7 note in LOJO stats para ────────────────────────────────
    for p in doc.paragraphs:
        if "ISU Four Continents 2022 used reduced panels)" in p.text:
            replace_in_para(p,
                "ISU Four Continents 2022 used reduced panels)",
                "ISU Four Continents 2022 Ice Dance segments used reduced panels; "
                "for J\u202f=\u202f7 events, LOJO replaces the removed judge with the median "
                "of the remaining six judges\u2019 marks)")
            print("  \u2713 E2b: J=7 median-of-6 note added")
            break
    else:
        print("  \u26a0 E2b: Four Continents text not found, skipping")

    # ── Table 5 changes (A6, A7, B2) ────────────────────────────────────────
    t5 = doc.tables[5]
    rows5 = [ch for ch in t5._tbl if ch.tag == f"{{{W}}}tr"]

    # A6 ISU names + A7 exact p-values
    # col indices: 0=Competition,1=Discipline,2=Segment,3=Judge,4=Bias,5=Margin,6=p,7=q,8=LOJO
    cell_updates = {
        1: {6: "0.0003"},
        2: {6: "0.0039"},   # OWG 2022 (was already 0.004, update to 4-decimal)
        3: {6: "0.0003"},
        4: {6: "0.0004"},
        5: {1: "Women Single Skating", 6: "0.0008"},
        6: {1: "Men Single Skating",   6: "0.0026"},  # GPF 23/24 (was 0.003)
        7: {6: "0.0006"},
        8: {1: "Men Single Skating",   6: "0.0009"},
        9: {1: "Women Single Skating", 6: "0.0003"},
    }
    for row_i, col_map in cell_updates.items():
        row = rows5[row_i]
        cells = [ch for ch in row if ch.tag == f"{{{W}}}tc"]
        for col_i, new_text in col_map.items():
            set_cell_text(cells[col_i], new_text)
    print("  \u2713 A6/A7: Table 5 discipline names and p-values updated")

    # B2: add LOJO CF Δ column
    cf_margins = ["0.25", "0.52", "1.62", "3.01", "4.26", "2.12", "1.57", "0.09", "5.72"]
    add_table_column(t5, "LOJO CF \u0394 (pts)", cf_margins)
    print("  \u2713 B2: LOJO CF \u0394 column added to Table 5")

    # Update Table 3 caption to explain new columns
    for p in doc.paragraphs:
        if "\u2020 Valieva results subsequently annulled by ISU; see Section 9.3." in p.text:
            replace_in_para(p,
                "\u2020 Valieva results subsequently annulled by ISU; see Section 9.3.",
                "\u2020 Valieva results subsequently annulled by ISU; see Section\u202f9.3. "
                "LOJO: whether removing the judge (replacing marks with the panel median) "
                "changes the gold-medal winner. LOJO\u202fCF\u202f\u0394: gold\u2013silver margin in the "
                "LOJO counterfactual (may involve different teams when LOJO\u202f=\u202fYes).")
            print("  \u2713 Table 3 caption: LOJO column explanations added")
            break
    else:
        print("  \u26a0 Table 3 caption: anchor text not found, skipping")

    # ── D2: Table B row label + note paragraph ───────────────────────────────
    t0 = doc.tables[0]
    rows0 = [ch for ch in t0._tbl if ch.tag == f"{{{W}}}tr"]
    for row in rows0:
        cells = [ch for ch in row if ch.tag == f"{{{W}}}tc"]
        txt = "".join(e.text or "" for e in cells[0].findall(f".//{{{W}}}t"))
        if "Panel judges (distinct roles)" in txt:
            set_cell_text(cells[0], "Panel judge-event positions")
            print("  \u2713 D2: Table B label updated")
            break

    for p in doc.paragraphs:
        if p.text.strip() == "Table B. Database characteristics.":
            insert_para_after(p,
                "Note: \u2018Panel judge-event positions\u2019 counts each judge\u2019s assignment "
                "to one event as one entry (one position per judge per event); the same "
                "individual may hold positions in multiple events across the database.")
            print("  \u2713 D2: Table B explanatory note added")
            break
    else:
        print("  \u26a0 D2: Table B caption para not found")

    # ── save ─────────────────────────────────────────────────────────────────
    doc.save(DOCX)
    print(f"\nSaved {DOCX}")

    # ── verification ─────────────────────────────────────────────────────────
    print("\n\u2500\u2500 Verification \u2500\u2500")
    doc2 = Document(DOCX)

    def tbl5_cell(ri, ci):
        rows = [ch for ch in doc2.tables[5]._tbl if ch.tag == f"{{{W}}}tr"]
        cells = [ch for ch in rows[ri] if ch.tag == f"{{{W}}}tc"]
        return "".join(e.text or "" for e in cells[ci].findall(f".//{{{W}}}t"))

    checks = [
        ("A1 'substantially attenuates'",
         lambda d: any("substantially attenuates" in p.text for p in d.paragraphs)),
        ("A2 old text GONE",
         lambda d: not any("not test validity" in p.text for p in d.paragraphs)),
        ("A2 new validity text",
         lambda d: any("exact conditional on the exchangeability" in p.text for p in d.paragraphs)),
        ("A3 §9.5 softened",
         lambda d: any("may reflect a mixture of true alternatives and residual departures" in p.text for p in d.paragraphs)),
        ("A4 §9.1 'consistent with a mixture'",
         lambda d: any("consistent with a mixture of true alternatives and null tests" in p.text for p in d.paragraphs)),
        ("A5 anonymization gone",
         lambda d: not any("anonymization complicates" in p.text for p in d.paragraphs)),
        ("A6a ISU names in §9.2",
         lambda d: any("Women Single Skating showed the highest proportion" in p.text for p in d.paragraphs)),
        ("A6b programme GONE",
         lambda d: not any("programme component" in p.text for p in d.paragraphs)),
        ("A7 exact p OWG2026",
         lambda _: tbl5_cell(1, 6) == "0.0003"),
        ("A7 exact p EUR2024",
         lambda _: tbl5_cell(5, 6) == "0.0008"),
        ("A8 J1 nationality softened",
         lambda d: any("We did not observe similarly unidirectional" in p.text for p in d.paragraphs)),
        ("B1 LOJO caution",
         lambda d: any("should not be interpreted as evidence of misconduct" in p.text for p in d.paragraphs)),
        ("B2 CF Δ column header",
         lambda _: "LOJO CF" in tbl5_cell(0, 9)),
        ("B2 CF Δ row 1 value",
         lambda _: tbl5_cell(1, 9) == "0.25"),
        ("B3 discordance in §9.4",
         lambda d: any("LOJO\u202f=\u202fNo rows" in p.text or "LOJO\u00a0=\u00a0No rows" in p.text
                       or "LOJO = No rows" in p.text or "LOJO\u202f=\u202fNo" in p.text
                       for p in d.paragraphs)),
        ("C1/C5/C6 future-work para",
         lambda d: any("Stratified residual-label permutation" in p.text for p in d.paragraphs)),
        ("C4 24-event calibration",
         lambda d: any("24 events (16.9%)" in p.text for p in d.paragraphs)),
        ("D2 Table B label",
         lambda d: any(
             "Panel judge-event positions" in "".join(e.text or "" for e in tc.findall(f".//{{{W}}}t"))
             for row in [ch for ch in d.tables[0]._tbl if ch.tag == f"{{{W}}}tr"]
             for tc in [ch for ch in row if ch.tag == f"{{{W}}}tc"])),
        ("E2 'remaining J−1'",
         lambda d: any("remaining J\u22121 judges" in p.text for p in d.paragraphs)),
    ]

    all_ok = True
    for label, fn in checks:
        ok = fn(doc2)
        sym = "\u2713" if ok else "\u2717"
        print(f"  {sym} {label}")
        if not ok:
            all_ok = False

    print("\nAll checks passed." if all_ok else "\nSome checks FAILED — review above.")


if __name__ == "__main__":
    main()
