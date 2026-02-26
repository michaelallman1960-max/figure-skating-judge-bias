"""
Converts six Judging Bias .md files to page-numbered .docx files.
Run with: python3 make_word_docs.py
"""

import os
import re
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/allman/Library/CloudStorage/Dropbox/Dropbox Mike/Judging Bias"

FILES = [
    ("journal_strategy.md",    "Journal & Publication Strategy"),
    ("media_strategy.md",      "Media & Op-Ed Strategy"),
    ("revenue_strategy.md",    "Revenue & Commercialization Strategy"),
    ("ip_protection.md",       "IP Protection & Legal Framework"),
    ("significance_draft_v1.md", "Significance Magazine — Draft v1"),
    ("oped_draft.md",          "Op-Ed Draft (NYT) — ISU-Impact Method"),
    ("cover_emails_2026-02-24.md", "Cover Emails — Full Package Submission"),
    ("methodology_diagnosis_v1.md", "Methodology Diagnosis — The Exchangeability Problem"),
    ("history_log.md",             "History Log — Session & Decision Record"),
    ("ISU_Scoring_Methodology.md", "ISU Ice Dance Scoring Methodology"),
    ("faq_v1.md",                  "FAQ — Statistical Audit of Figure Skating Judging"),
    ("Data_Dictionary.md",         "Data Dictionary — ISU Figure Skating IJS Judging Bias Study"),
]

TODAY = date.today().strftime("%B %d, %Y")


def add_page_number(paragraph):
    """Insert a PAGE field into the paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' PAGE '
    run._r.append(instr)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)


def add_num_pages(paragraph):
    """Insert a NUMPAGES field into the paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_begin)

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = ' NUMPAGES '
    run._r.append(instr)

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_end)


def setup_footer(doc, short_title):
    """Add a footer with short title (left), page X of Y (right)."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear existing footer content
    for para in footer.paragraphs:
        for run in para.runs:
            run.text = ""

    ftable = footer.add_table(1, 2, width=Inches(6.5))
    ftable.style = 'Table Grid'
    # Remove table borders
    tbl = ftable._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.append(tblPr)

    # Left cell: short title
    left_cell = ftable.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_para.add_run(f"Allman — {short_title}")
    left_run.font.size = Pt(9)
    left_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # Right cell: page X of Y
    right_cell = ftable.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = right_para.add_run("Page ")
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    add_page_number(right_para)
    r2 = right_para.add_run(" of ")
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    add_num_pages(right_para)


def parse_md_line(line):
    """
    Returns (kind, text, level) where kind is one of:
    'h1', 'h2', 'h3', 'hr', 'bullet', 'code', 'blank', 'table_row', 'italic_line', 'text'
    """
    stripped = line.rstrip()

    if stripped.startswith('# '):
        return ('h1', stripped[2:].strip(), 1)
    if stripped.startswith('## '):
        return ('h2', stripped[3:].strip(), 2)
    if stripped.startswith('### '):
        return ('h3', stripped[4:].strip(), 3)
    if stripped.startswith('#### '):
        return ('h3', stripped[5:].strip(), 4)
    if re.match(r'^-{3,}$', stripped) or re.match(r'^\*{3,}$', stripped):
        return ('hr', '', 0)
    if stripped.startswith('- ') or stripped.startswith('* '):
        return ('bullet', stripped[2:].strip(), 0)
    if re.match(r'^\d+\. ', stripped):
        return ('numbered', re.sub(r'^\d+\. ', '', stripped), 0)
    if stripped.startswith('|'):
        return ('table_row', stripped, 0)
    if stripped.startswith('```'):
        return ('code_fence', stripped, 0)
    if (stripped.startswith('*') and stripped.endswith('*') and not stripped.startswith('**')):
        return ('italic_line', stripped.strip('*'), 0)
    if stripped == '':
        return ('blank', '', 0)
    return ('text', stripped, 0)


def strip_inline_md(text):
    """Remove bold/italic/code markers from inline text for plain Word runs."""
    # We handle bold/italic via add_formatted_run instead
    return text


def add_formatted_run(paragraph, text, bold=False, italic=False, size=None, color=None):
    """Add a run with formatting."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_paragraph_with_inline(doc_or_para, text, style=None, bold_base=False, size=None):
    """
    Parse inline markdown (bold, italic, code) and add to a new paragraph.
    Returns the paragraph.
    """
    if style:
        para = doc_or_para.add_paragraph(style=style)
    else:
        para = doc_or_para.add_paragraph()

    # Parse inline: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*|`.*?`|~~.*?~~)')
    parts = pattern.split(text)

    for part in parts:
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            r = para.add_run(part[3:-3])
            r.bold = True
            r.italic = True
        elif part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = para.add_run(part[1:-1])
            r.italic = True
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
        else:
            r = para.add_run(part)
            if bold_base:
                r.bold = True
        if size:
            r.font.size = Pt(size)

    return para


def md_to_docx(md_path, out_path, doc_title, short_title):
    doc = Document()

    # Page setup: 1-inch margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title page block
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = title_para.add_run(doc_title)
    tr.bold = True
    tr.font.size = Pt(16)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sr = sub_para.add_run(f"Michael Allman  |  {TODAY}  |  Judging Bias in Competitive Figure Skating")
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()  # spacer

    # Setup footer
    setup_footer(doc, short_title)

    # Read markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    skip_first_h1 = True  # The file's own H1 becomes the doc title above

    i = 0
    while i < len(lines):
        raw = lines[i]
        kind, text, level = parse_md_line(raw)

        # Code block handling
        if kind == 'code_fence':
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # End of code block — dump as monospace
                if code_lines:
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run('\n'.join(code_lines))
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(9)
                    code_para.paragraph_format.left_indent = Inches(0.25)
                in_code_block = False
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(raw.rstrip())
            i += 1
            continue

        # Table handling
        if kind == 'table_row':
            # Skip separator rows (|---|---|)
            if re.match(r'^\|[\s\-\|:]+\|$', text.strip()):
                i += 1
                continue
            cells = [c.strip() for c in text.strip('|').split('|')]
            if not in_table:
                # First row = header
                in_table = True
                num_cols = len(cells)
                table = doc.add_table(rows=1, cols=num_cols)
                table.style = 'Table Grid'
                hdr_row = table.rows[0]
                for ci, cell_text in enumerate(cells):
                    cell = hdr_row.cells[ci]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    r = p.add_run(re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text))
                    r.bold = True
                    r.font.size = Pt(10)
            else:
                row_cells = table.add_row().cells
                for ci, cell_text in enumerate(cells[:len(row_cells)]):
                    row_cells[ci].text = ''
                    p = row_cells[ci].paragraphs[0]
                    # Inline bold in cells
                    add_formatted_run(p, re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text), size=10)
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                doc.add_paragraph()  # space after table

        # Normal line kinds
        if kind == 'h1':
            if skip_first_h1:
                skip_first_h1 = False
                i += 1
                continue
            p = doc.add_paragraph(style='Heading 1')
            p.clear()
            add_formatted_run(p, text, bold=True, size=14)
        elif kind == 'h2':
            p = doc.add_paragraph(style='Heading 2')
            p.clear()
            add_formatted_run(p, text, bold=True, size=13)
        elif kind == 'h3':
            p = doc.add_paragraph(style='Heading 3')
            p.clear()
            add_formatted_run(p, text, bold=True, size=12)
        elif kind == 'hr':
            # Horizontal rule — add a line via border on paragraph
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'AAAAAA')
            pBdr.append(bottom)
            pPr.append(pBdr)
        elif kind == 'bullet':
            p = add_paragraph_with_inline(doc, text, style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
        elif kind == 'numbered':
            p = add_paragraph_with_inline(doc, text, style='List Number')
            p.paragraph_format.left_indent = Inches(0.25)
        elif kind == 'italic_line':
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.italic = True
            r.font.size = Pt(11)
        elif kind == 'blank':
            # Only add spacing if previous paragraph wasn't already blank
            pass  # Word's paragraph spacing handles this
        elif kind == 'text':
            add_paragraph_with_inline(doc, text)

        i += 1

    doc.save(out_path)
    print(f"  ✓ Saved: {os.path.basename(out_path)}")


def main():
    print(f"\nGenerating Word documents — {TODAY}\n")
    for md_file, title in FILES:
        md_path = os.path.join(BASE, md_file)
        out_name = md_file.replace('.md', '.docx')
        out_path = os.path.join(BASE, out_name)
        short = title.split('—')[0].strip().split('(')[0].strip()
        if not os.path.exists(md_path):
            print(f"  ✗ NOT FOUND: {md_file}")
            continue
        print(f"  Converting: {md_file}")
        try:
            md_to_docx(md_path, out_path, title, short)
        except Exception as e:
            print(f"  ✗ ERROR on {md_file}: {e}")
    print("\nDone.\n")


if __name__ == '__main__':
    main()
