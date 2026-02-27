"""
fix_eighth_review.py
Eighth-review fixes (items from ChatGPT review 8):

  1. LOJO arithmetic / "all 142 have 9 judges" — fix §9 methodological notes
  2. p-value support notation: {1/M,...,1} → {1/(M+1),...,1}  (§9.5)
  3. "7.4% discovery rate conservative relative to 5%" — rewrite §9.5
  4. Table A last row — add calibration qualifier
  5. "at most 0.01 pts" — soften to "on the order of hundredths of a point"  (§7)
  6. "judge-competitor pairings" → "judge × competitor-pair results"  (Abstract + Conclusion)
  7. "hierarchical FDR (Benjamini–Yekutieli)" → "dependence-robust FDR (Benjamini–Yekutieli)"  (§10)
  8. Table 3 caption — clarify LOJO CF Δ definition
"""

import shutil
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX = Path("judge_bias_isu_judging_system.docx")
BAK  = Path("judge_bias_isu_judging_system.docx.bak_eighth_review")
W    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def replace_in_para(para, old, new):
    """Replace old→new in paragraph. Zeroes all runs, puts text in run 0."""
    full = para.text
    if old not in full:
        raise ValueError(f"Text not found: {repr(old[:80])}")
    new_text = full.replace(old, new, 1)
    for i, run in enumerate(para.runs):
        run.text = new_text if i == 0 else ""


def replace_in_table_cell(table, search_text, old, new):
    """Find a cell containing search_text in table and replace old→new in it."""
    W_ns = f"{{{W}}}"
    for row in table.rows:
        for cell in row.cells:
            if search_text in cell.text:
                # Find the paragraph containing the text
                for para in cell.paragraphs:
                    if old in para.text:
                        full = para.text
                        new_text = full.replace(old, new, 1)
                        for i, run in enumerate(para.runs):
                            run.text = new_text if i == 0 else ""
                        return True
    return False


def main():
    shutil.copy(DOCX, BAK)
    print(f"Backed up to {BAK}")

    doc = Document(DOCX)

    # ── Fix 1: Remove "All 142 analyzed events have panels of exactly nine judges." ──
    # Replace with accurate statement about 4CC 2022 eight-judge events
    OLD_1 = (
        "Events with fewer than nine judges are excluded from permutation inference "
        "(n=2 events). All 142 analyzed events have panels of exactly nine judges. "
        "LOJO is computed for all 144 events, including the two events with reduced "
        "panels (seven judges each)."
    )
    NEW_1 = (
        "Events with fewer than nine judges are excluded from permutation inference "
        "(n=2 events). Of the 142 analyzed events, 138 have standard nine-judge panels; "
        "the four remaining events (ISU Four Continents Championships 2022 Men Single "
        "Skating and Pair Skating segments) used eight-judge panels. "
        "LOJO is computed for all 144 events using the actual panel size for each event."
    )
    found = False
    for p in doc.paragraphs:
        if OLD_1 in p.text:
            replace_in_para(p, OLD_1, NEW_1)
            print("  \u2713 Fix 1: eight-judge panel caveat added to \u00a79 methodological notes")
            found = True
            break
    if not found:
        raise ValueError("Fix 1: methodological notes text not found")

    # ── Fix 2: p-value support notation {1/M,...} → {1/(M+1),...} ──────────────
    OLD_2 = (
        "under the null, the permutation p-value is discrete-uniform on "
        "{1/M, 2/M, \u2026, 1} (converging to continuous uniform(0,\u202f1) as M\u2192\u221e)"
    )
    NEW_2 = (
        "under the null, the permutation p-value is discrete-uniform on "
        "{1/(M+1), 2/(M+1), \u2026, 1} (converging to continuous uniform(0,\u202f1) as M\u2192\u221e)"
    )
    found = False
    for p in doc.paragraphs:
        if OLD_2 in p.text:
            replace_in_para(p, OLD_2, NEW_2)
            print("  \u2713 Fix 2: p-value support corrected to {1/(M+1),...,1}")
            found = True
            break
    if not found:
        raise ValueError("Fix 2: p-value support text not found")

    # ── Fix 3: "7.4% discovery rate conservative relative to 5.0%" ──────────────
    OLD_3 = (
        "The 7.4% BH-FDR discovery rate (q\u2264\u200a0.05) is conservative relative to "
        "the nominal 5.0% by design, consistent with a non-trivial fraction of "
        "true alternatives."
    )
    NEW_3 = (
        "BH controls the expected false discovery proportion among the rejected "
        "hypotheses at 5% within each event; the 7.4% rejection proportion "
        "indicates a substantial fraction of tests show competitor-specific "
        "differentials under this procedure."
    )
    found = False
    for p in doc.paragraphs:
        if OLD_3 in p.text:
            replace_in_para(p, OLD_3, NEW_3)
            print("  \u2713 Fix 3: '7.4% conservative relative to 5%' sentence rewritten")
            found = True
            break
    if not found:
        raise ValueError("Fix 3: '7.4% conservative' text not found")

    # ── Fix 4: Table A last row — add calibration qualifier ─────────────────────
    # Table A is doc.tables[1] (index 1: "Alternative null models evaluated")
    # The last row Status cell says "Adopted ✓" with text about no violations
    OLD_4 = (
        "No violations identified in diagnostics to date; exchangeability is plausible "
        "given median-of-others neutralization removes the shared competitor-quality signal"
    )
    NEW_4 = (
        "No violations identified in diagnostics to date; exchangeability is plausible "
        "given median-of-others neutralization removes the shared competitor-quality signal. "
        "Exchangeability remains an assumption and is addressed via planned calibration checks."
    )
    found = False
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                full_text = cell.text
                if OLD_4 in full_text:
                    for para in cell.paragraphs:
                        if OLD_4 in para.text:
                            replace_in_para(para, OLD_4, NEW_4)
                            found = True
                            break
                if found:
                    break
            if found:
                break
        if found:
            break
    if found:
        print("  \u2713 Fix 4: Table A calibration qualifier added")
    else:
        raise ValueError("Fix 4: Table A 'No violations identified' text not found")

    # ── Fix 5: "at most 0.01 pts" → soften rounding claim ───────────────────────
    OLD_5 = (
        "the bias B_j is computed on raw unrounded marks, so differences are at most 0.01 pts"
    )
    NEW_5 = (
        "the bias B_j is computed on raw unrounded marks, so differences are small "
        "(on the order of hundredths of a point, due to ISU rounding to two decimal places)"
    )
    found = False
    for p in doc.paragraphs:
        if OLD_5 in p.text:
            replace_in_para(p, OLD_5, NEW_5)
            print("  \u2713 Fix 5: rounding claim softened in \u00a77")
            found = True
            break
    if not found:
        raise ValueError("Fix 5: rounding claim text not found")

    # ── Fix 6: "judge-competitor pairings" → "judge \u00d7 competitor-pair results" ──
    # In Abstract and Conclusion
    OLD_6 = "20,213 significant judge-competitor pairings"
    NEW_6 = "20,213 significant judge \u00d7 competitor-pair results"
    count = 0
    for p in doc.paragraphs:
        if OLD_6 in p.text:
            replace_in_para(p, OLD_6, NEW_6)
            count += 1
    if count == 0:
        raise ValueError("Fix 6: 'judge-competitor pairings' not found")
    print(f"  \u2713 Fix 6: 'judge-competitor pairings' \u2192 'judge \u00d7 competitor-pair results' ({count} occurrence(s))")

    # ── Fix 7: "hierarchical FDR (Benjamini\u2013Yekutieli)" → "dependence-robust FDR" ──
    OLD_7 = "hierarchical FDR (Benjamini\u2013Yekutieli)"
    NEW_7 = "dependence-robust FDR (Benjamini\u2013Yekutieli)"
    found = False
    for p in doc.paragraphs:
        if OLD_7 in p.text:
            replace_in_para(p, OLD_7, NEW_7)
            print("  \u2713 Fix 7: 'hierarchical FDR' \u2192 'dependence-robust FDR'")
            found = True
            break
    if not found:
        raise ValueError("Fix 7: 'hierarchical FDR (Benjamini\u2013Yekutieli)' not found")

    # ── Fix 8: Table 3 caption — clarify LOJO CF \u0394 definition ──────────────────
    OLD_8 = (
        "LOJO\u202fCF\u202f\u0394: gold\u2013silver margin in the LOJO counterfactual "
        "(may involve different teams when LOJO\u202f=\u202fYes)."
    )
    NEW_8 = (
        "LOJO\u202fCF\u202f\u0394: the margin between the top-two finishers in the LOJO "
        "counterfactual standings (which may be a different pair of teams than the "
        "original gold\u2013silver when LOJO\u202f=\u202fYes; CF\u202f\u0394 is therefore "
        "not directly comparable to the Margin\u202f(pts) column in LOJO\u202f=\u202fYes rows)."
    )
    found = False
    for p in doc.paragraphs:
        if OLD_8 in p.text:
            replace_in_para(p, OLD_8, NEW_8)
            print("  \u2713 Fix 8: Table 3 caption LOJO CF \u0394 definition clarified")
            found = True
            break
    if not found:
        # Also check in table captions (sometimes in table rows)
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if OLD_8 in cell.text:
                        for para in cell.paragraphs:
                            if OLD_8 in para.text:
                                replace_in_para(para, OLD_8, NEW_8)
                                found = True
                                break
                        if found:
                            break
                if found:
                    break
            if found:
                break
    if not found:
        raise ValueError("Fix 8: Table 3 LOJO CF \u0394 caption text not found")

    doc.save(DOCX)
    print(f"\nSaved {DOCX}")

    # ── Verification ─────────────────────────────────────────────────────────────
    print("\n\u2500\u2500 Verification \u2500\u2500")
    doc2 = Document(DOCX)

    checks = [
        ("Fix 1: eight-judge caveat",
         lambda d: any("138 have standard nine-judge panels" in p.text for p in d.paragraphs)),
        ("Fix 1: old 'All 142' GONE",
         lambda d: not any("All 142 analyzed events have panels of exactly nine judges" in p.text for p in d.paragraphs)),
        ("Fix 2: {1/(M+1),...}",
         lambda d: any("1/(M+1), 2/(M+1)" in p.text for p in d.paragraphs)),
        ("Fix 2: old {1/M,...} GONE",
         lambda d: not any("{1/M, 2/M" in p.text for p in d.paragraphs)),
        ("Fix 3: 'rejection proportion'",
         lambda d: any("7.4% rejection proportion" in p.text for p in d.paragraphs)),
        ("Fix 4: calibration qualifier in Table A",
         lambda d: any(
             "Exchangeability remains an assumption" in cell.text
             for tbl in d.tables for row in tbl.rows for cell in row.cells
         )),
        ("Fix 5: rounding softened",
         lambda d: any("on the order of hundredths of a point" in p.text for p in d.paragraphs)),
        ("Fix 6: judge x pair results",
         lambda d: any("judge \u00d7 competitor-pair results" in p.text for p in d.paragraphs)),
        ("Fix 7: dependence-robust FDR",
         lambda d: any("dependence-robust FDR (Benjamini\u2013Yekutieli)" in p.text for p in d.paragraphs)),
        ("Fix 8: CF \u0394 clarified",
         lambda d: any("not directly comparable to the Margin" in p.text for p in d.paragraphs)),
    ]

    all_ok = True
    for label, check_fn in checks:
        ok = check_fn(doc2)
        symbol = "\u2713" if ok else "\u2717"
        print(f"  {symbol} {label}")
        if not ok:
            all_ok = False

    if all_ok:
        print("\nAll 8 fixes verified.")
    else:
        print("\nSome checks FAILED \u2014 review above.")


if __name__ == "__main__":
    main()
