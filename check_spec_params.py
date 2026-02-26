#!/usr/bin/env python3
"""
check_spec_params.py
====================
Verifies that the engineering spec and reproduction checklist match the
actual method parameters stored in the database. Run before any submission
commit to catch spec drift early.

Usage:
    python3 check_spec_params.py

Exit codes:
    0 — all checks passed
    1 — one or more checks failed
"""

import sqlite3
import sys
from pathlib import Path

BASE = Path(__file__).parent
DB_PATH = BASE / "figure_skating_ijs_v4.sqlite"
ENG_SPEC = BASE / "engineering_spec_isuimpact_v1.docx"
REPRO_CHECKLIST = BASE / "reproduction_checklist_isuimpact.docx"

# ── Expected values (ground truth from the published run) ─────────────────────
EXPECTED = {
    "method_version": "isuimpact_residual_v1",
    "rng_seed": 20260223,
    "permutations": 10000,
    "events_analyzed": 142,
    "cdf_scope": "global",
}

PASS = "  ✅"
FAIL = "  ❌"

failures = []


def check(label, ok, detail=""):
    status = PASS if ok else FAIL
    print(f"{status}  {label}")
    if not ok:
        print(f"       {detail}")
        failures.append(label)


# ── 1. Database checks ─────────────────────────────────────────────────────────
print("\n── Database (figure_skating_ijs_v4.sqlite) ──────────────────────────────")

if not DB_PATH.exists():
    print(f"{FAIL}  Database not found: {DB_PATH}")
    failures.append("database_exists")
else:
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()

    # Method version — verify the published method is present
    rows = cur.execute(
        "SELECT DISTINCT method_version FROM pairwise_impact_results"
    ).fetchall()
    versions = [r[0] for r in rows]
    check(
        f"method_version '{EXPECTED['method_version']}' present",
        EXPECTED["method_version"] in versions,
        f"Found: {versions}",
    )

    # RNG seed
    seeds = cur.execute(
        "SELECT DISTINCT rng_seed FROM pairwise_impact_results"
    ).fetchall()
    seeds = [r[0] for r in seeds]
    check(
        f"rng_seed = {EXPECTED['rng_seed']}",
        seeds == [EXPECTED["rng_seed"]],
        f"Found: {seeds}",
    )

    # Permutations
    perms = cur.execute(
        "SELECT DISTINCT permutations FROM pairwise_impact_results"
    ).fetchall()
    perms = [r[0] for r in perms]
    check(
        f"permutations = {EXPECTED['permutations']} (flat, all events)",
        perms == [EXPECTED["permutations"]],
        f"Found: {perms}",
    )

    # Events analyzed
    n_events = cur.execute(
        "SELECT COUNT(DISTINCT event_id) FROM pairwise_impact_results"
    ).fetchone()[0]
    check(
        f"events_analyzed = {EXPECTED['events_analyzed']}",
        n_events == EXPECTED["events_analyzed"],
        f"Found: {n_events}",
    )

    # Row counts sanity — check published method version specifically
    n_rows_v2 = cur.execute(
        "SELECT COUNT(*) FROM pairwise_impact_results WHERE method_version=?",
        (EXPECTED["method_version"],)
    ).fetchone()[0]
    check(
        f"pairwise_impact_results rows for {EXPECTED['method_version']} = 271,728",
        n_rows_v2 == 271728,
        f"Found: {n_rows_v2:,}",
    )

    conn.close()


# ── 2. Engineering spec checks ─────────────────────────────────────────────────
print("\n── Engineering Spec (engineering_spec_isuimpact_v1.docx) ────────────────")

try:
    from docx import Document

    def doc_text(path):
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    spec_text = doc_text(ENG_SPEC)

    check(
        "Seed 20260223 mentioned",
        "20260223" in spec_text,
        "String '20260223' not found in spec",
    )
    check(
        "M=10,000 mentioned",
        "10,000" in spec_text or "10000" in spec_text,
        "No mention of 10,000 permutations",
    )
    check(
        "Residual-label method described",
        "residual" in spec_text.lower() or "delta-exchange" in spec_text.lower(),
        "No mention of residual-label / delta-exchange method in spec",
    )
    check(
        "No quantile/CDF permutation language",
        "percentile labels" not in spec_text
        and "empirical CDF" not in spec_text.lower(),
        "Found stale quantile-null language (percentile labels / empirical CDF)",
    )
    check(
        "goe_factors table referenced",
        "goe_factors" in spec_text,
        "No mention of goe_factors reference table",
    )
    check(
        "No adaptive permutation strategy",
        "M=1,000" not in spec_text and "M=1000" not in spec_text,
        "Found stale adaptive strategy reference (M=1,000)",
    )
    check(
        "version_stamp present (Version 1.1 or later)",
        any(f"Version 1.{n}" in spec_text for n in range(1, 10))
        or "Version 2" in spec_text,
        "Spec still at Version 1.0 — needs update",
    )

except ImportError:
    print(f"{FAIL}  python-docx not installed; skipping spec text checks")
    failures.append("docx_import")


# ── 3. Reproduction checklist checks ──────────────────────────────────────────
print("\n── Reproduction Checklist (reproduction_checklist_isuimpact.docx) ───────")

try:
    checklist_text = doc_text(REPRO_CHECKLIST)

    check(
        "Seed 20260223 mentioned",
        "20260223" in checklist_text,
        "String '20260223' not found in checklist",
    )
    check(
        "M=10,000 mentioned",
        "10,000" in checklist_text or "Nperm = 10" in checklist_text
        or "10000" in checklist_text,
        "No mention of 10,000 permutations",
    )
    check(
        "Global CDF scope stated",
        "global" in checklist_text.lower(),
        "No mention of 'global' CDF scope — replicator will use event-scope",
    )
    check(
        "No event-scope-only CDF language",
        "across the event" not in checklist_text
        and "for that event" not in checklist_text,
        "Found stale event-scope language",
    )
    check(
        "Spec audit step present",
        "check_spec_params" in checklist_text
        or "spec audit" in checklist_text.lower()
        or "engineering spec" in checklist_text.lower(),
        "No spec audit step found in checklist",
    )

except Exception as e:
    print(f"{FAIL}  Could not read checklist: {e}")
    failures.append("checklist_read")


# ── Summary ────────────────────────────────────────────────────────────────────
print()
if failures:
    print(f"❌  FAILED — {len(failures)} check(s) failed:")
    for f in failures:
        print(f"     • {f}")
    print()
    sys.exit(1)
else:
    print("✅  ALL CHECKS PASSED — spec, checklist, and database are in sync.")
    print()
    sys.exit(0)
