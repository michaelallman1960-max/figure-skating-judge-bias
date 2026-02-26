"""
fix_section6_cleanup.py
Seven structural fixes based on second ChatGPT review:

1. Para [21]: Strip stale 6.2/6.3/6.4 (quantile null labeled 'primary') from 6.1 paragraph
2. Para [30]: Fix OWG null attribution ('style-adjusted distribution-preserving' → 'residual-label')
3. Para [79]: Appendix A — tag 17.78% as v1-only, add v2 contrast, soften 'Type I inflation'
4. Table 4 (paper Table 2): Add footnote clarifying team-events exclusion vs Section 9.1
5. Para [39]: Remove 'CDF scope is global' sentence (describes v1, not the adopted v2 method)
6. Table 1 (paper Table A): Soften 'exchangeability holds by construction' claim
7. Para [31]: Fix Table 1 caption null name ('style-adjusted' → 'residual-label')
"""

import shutil
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX_PATH = Path("judge_bias_isu_judging_system.docx")
BACKUP_PATH = Path("judge_bias_isu_judging_system.docx.bak_section6_cleanup")

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def backup():
    shutil.copy2(DOCX_PATH, BACKUP_PATH)
    print(f"Backup → {BACKUP_PATH}")


def replace_in_para(para, old, new, label=""):
    full = para.text
    if old not in full:
        raise ValueError(f"[{label}] text not found:\n  '{old[:100]}'")
    new_text = full.replace(old, new)
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def find_para(doc, substring):
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    raise ValueError(f"Paragraph containing '{substring[:60]}' not found")


# ──────────────────────────────────────────────────────────────
# FIX 1: Para [21] — strip stale 6.2/6.3/6.4 quantile null content
# Keep only the 6.1 rationale sentence.
# ──────────────────────────────────────────────────────────────

def fix_para21_strip_old_quantile(doc):
    p21 = find_para(doc, "6.2. Distribution-preserving style-adjusted null (primary)")
    old_block = (
        "\n\n6.2. Distribution-preserving style-adjusted null (primary). "
        "Partition rows into scoring categories c (GOE and each PCS component). "
        "For each judge j and category c, form the empirical CDF F_{j,c} "
        "from all marks that judge has recorded across all events in the database "
        "(global CDF scope); using the full scoring history yields a more stable "
        "style estimate than a single event\u2019s marks alone. Map each observed mark "
        "to a percentile u_j,r,T = F_j,c(x_j,r,T), using randomized tie-breaking "
        "to handle discreteness. Under the null, within each row we permute the nine "
        "percentiles across judge labels, then map percentiles back to marks using the "
        "inverse empirical CDF of the receiving judge: x_perm_j,r,T = "
        "F^{-1}_j,c(u_perm_j,r,T). This preserves each judge\u2019s full marginal "
        "distribution of marks within each category while breaking any stable "
        "judge\u2013competitor association."
        "\n\n6.3. Permutation p-values. For each (j,A,B), compute the observed "
        "B_j(A,B). Generate M permutations under the null, recompute B_perm_j(A,B) "
        "each time, and assign a two-sided Monte Carlo p-value:\n"
        "    p = (1 + count(|B_perm| >= |B_obs|)) / (M + 1)."
        "\n\n6.4. Multiple comparisons. In a segment with N competitors and J judges "
        "there are J*N*(N-1)/2 pairwise tests. We report Benjamini\u2013Hochberg false "
        "discovery rate (FDR) adjusted q-values within each event segment, and we "
        "discuss hierarchical control strategies for multi-event studies."
    )
    replace_in_para(p21, old_block, "", "fix1 strip old 6.2-6.4")
    print("  Stripped stale 6.2/6.3/6.4 quantile null block from para [21].")


# ──────────────────────────────────────────────────────────────
# FIX 2: Para [30] — fix OWG null attribution
# ──────────────────────────────────────────────────────────────

def fix_owg_null_attribution(doc):
    p30 = find_para(doc, "style-adjusted distribution-preserving null, judge J1 shifts")
    replace_in_para(
        p30,
        "Under the style-adjusted distribution-preserving null, judge J1 shifts",
        "Under the residual-label (delta-exchange) permutation null, judge J1 shifts",
        "fix2 OWG null name"
    )
    print("  Fixed OWG paragraph null attribution: 'style-adjusted distribution-preserving' → 'residual-label'.")


# ──────────────────────────────────────────────────────────────
# FIX 3: Para [79] Appendix A — tag 17.78% as v1-only, add v2 contrast
# ──────────────────────────────────────────────────────────────

def fix_appendix_a_v1_tag(doc):
    p79 = find_para(doc, "Empirically, the quantile null yields nominal p")
    old_sent = (
        "Empirically, the quantile null yields nominal p \u2264 0.05 for 17.78% of "
        "pairwise tests versus the 5.0% expected under a well-calibrated null "
        "(3.6\u00d7 inflation), indicating systematic rejection of true-null cases."
    )
    new_sent = (
        "Empirically (v1 diagnostic only), the quantile null yields a nominal "
        "rejection rate of 17.78% at p \u2264 0.05 versus the 5.0% null expectation "
        "\u2014 far above the expected level under a well-calibrated test. By contrast, "
        "the adopted v2 residual-label method yields 25.4% nominal p \u2264 0.05 across "
        "the database (Section\u00a09.1), consistent with a null/alternative mixture "
        "rather than miscalibration."
    )
    replace_in_para(p79, old_sent, new_sent, "fix3 appendix A v1 tag")
    print("  Updated Appendix A: tagged 17.78% as v1-only, added v2 contrast.")


# ──────────────────────────────────────────────────────────────
# FIX 4: Table 2 footnote — clarify team-events exclusion
# Inserts text into the empty paragraph after Table 4 (paper's Table 2)
# ──────────────────────────────────────────────────────────────

def fix_table2_footnote(doc):
    body = doc.element.body
    children = list(body)
    table2_body_idx = None
    for i, child in enumerate(children):
        tag = child.tag.split("}")[1] if "}" in child.tag else child.tag
        if tag == "tbl":
            rows = child.findall(f".//{{{W}}}tr")
            if rows:
                first_texts = [e.text or "" for e in rows[0].findall(f".//{{{W}}}t")]
                if first_texts and first_texts[0].strip() == "Discipline":
                    table2_body_idx = i
                    break

    assert table2_body_idx is not None, "Could not find Table 2 (Discipline) in body"

    # The paragraph immediately after Table 2 (body[table2_body_idx + 1])
    after_elem = children[table2_body_idx + 1]
    tag = after_elem.tag.split("}")[1] if "}" in after_elem.tag else after_elem.tag
    assert tag == "p", f"Expected paragraph after Table 2, got {tag}"

    # Set text in that paragraph
    # Clear existing runs
    for r in after_elem.findall(f".//{{{W}}}r"):
        after_elem.remove(r)

    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = (
        "Note: Totals in this table exclude team event segments (8 events, 1,899 additional "
        "tests, 41 additional significant pairs). Full-database totals including team segments "
        "are reported in Section\u00a09.1 (271,728 tests; 20,213 significant pairs)."
    )
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r.append(t)
    after_elem.append(r)
    print("  Added footnote to Table 2 clarifying team-events exclusion vs Section 9.1.")


# ──────────────────────────────────────────────────────────────
# FIX 5: Para [39] — remove 'CDF scope is global' sentence
# ──────────────────────────────────────────────────────────────

def fix_remove_cdf_scope(doc):
    p39 = find_para(doc, "CDF scope is global")
    old_sent = (
        " CDF scope is global: each judge's style distribution is estimated from "
        "all their marks in the database, with event-local fallback for judges new to "
        "the database."
    )
    replace_in_para(p39, old_sent, "", "fix5 remove CDF scope")
    print("  Removed 'CDF scope is global' sentence from Methodological notes.")


# ──────────────────────────────────────────────────────────────
# FIX 6: Table A — soften 'exchangeability holds by construction'
# ──────────────────────────────────────────────────────────────

def fix_table_a_exchangeability(doc):
    old_claim = (
        "None identified; exchangeability holds by construction after removing "
        "the shared quality signal"
    )
    new_claim = (
        "No violations identified in diagnostics to date; exchangeability is plausible "
        "given median-of-others neutralization removes the shared competitor-quality signal"
    )

    found = False
    for t in doc.tables:
        tbl = t._tbl
        rows = tbl.findall(f".//{{{W}}}tr")
        if not rows:
            continue
        first_texts = [e.text or "" for e in rows[0].findall(f".//{{{W}}}t")]
        if "Method" not in str(first_texts):
            continue
        # Found Table A — scan all cells
        for row in rows:
            for tc in row.findall(f".//{{{W}}}tc"):
                cell_text = "".join(e.text or "" for e in tc.findall(f".//{{{W}}}t"))
                if old_claim in cell_text:
                    # Replace in each t element
                    for t_elem in tc.findall(f".//{{{W}}}t"):
                        if t_elem.text and old_claim in t_elem.text:
                            t_elem.text = t_elem.text.replace(old_claim, new_claim)
                            found = True
                    # Also try across multiple t elements
                    if not found:
                        # Rebuild: collect all t texts, replace, set first t
                        t_elems = tc.findall(f".//{{{W}}}t")
                        combined = "".join(e.text or "" for e in t_elems)
                        if old_claim in combined:
                            new_combined = combined.replace(old_claim, new_claim)
                            for e in t_elems:
                                e.text = ""
                            if t_elems:
                                t_elems[0].text = new_combined
                            found = True
        if found:
            break

    assert found, "Table A exchangeability claim not found"
    print("  Softened Table A 'exchangeability holds by construction' → diagnostic language.")


# ──────────────────────────────────────────────────────────────
# FIX 7: Para [31] Table 1 caption — fix null name
# ──────────────────────────────────────────────────────────────

def fix_table1_caption(doc):
    p31 = find_para(doc, "style-adjusted permutation null")
    replace_in_para(
        p31,
        "style-adjusted permutation null",
        "residual-label permutation null",
        "fix7 table1 caption"
    )
    print("  Fixed Table 1 caption: 'style-adjusted permutation null' → 'residual-label permutation null'.")


# ──────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────

def main():
    backup()
    doc = Document(str(DOCX_PATH))

    print("\n--- Fix 1: Strip stale 6.2/6.3/6.4 quantile null from para [21] ---")
    fix_para21_strip_old_quantile(doc)

    print("\n--- Fix 2: OWG paragraph null attribution ---")
    fix_owg_null_attribution(doc)

    print("\n--- Fix 3: Appendix A v1 tag + v2 contrast ---")
    fix_appendix_a_v1_tag(doc)

    print("\n--- Fix 4: Table 2 team-events footnote ---")
    fix_table2_footnote(doc)

    print("\n--- Fix 5: Remove CDF scope sentence from Methodological notes ---")
    fix_remove_cdf_scope(doc)

    print("\n--- Fix 6: Soften Table A exchangeability claim ---")
    fix_table_a_exchangeability(doc)

    print("\n--- Fix 7: Fix Table 1 caption null name ---")
    fix_table1_caption(doc)

    doc.save(str(DOCX_PATH))
    print(f"\nSaved → {DOCX_PATH}")

    # Verification
    doc2 = Document(str(DOCX_PATH))
    checks_present = [
        "6.1. Why style control is required",
        "residual-label (delta-exchange) permutation null, judge J1 shifts",
        "v1 diagnostic only",
        "25.4% nominal p",
        "1,899 additional tests",
        "residual-label permutation null",
        "No violations identified in diagnostics to date",
    ]
    checks_absent = [
        "Distribution-preserving style-adjusted null (primary)",
        "style-adjusted distribution-preserving null, judge J1",
        "style-adjusted permutation null",
        "CDF scope is global",
        "exchangeability holds by construction",
        "6.4. Multiple comparisons",
    ]
    print("\nVerification:")
    all_ok = True
    for text in checks_present:
        found = any(text in p.text for p in doc2.paragraphs)
        status = "OK" if found else "FAIL"
        if not found:
            all_ok = False
        print(f"  [{status}] present: '{text[:70]}'")
    for text in checks_absent:
        found = any(text in p.text for p in doc2.paragraphs)
        status = "OK" if not found else "FAIL"
        if found:
            all_ok = False
        print(f"  [{status}] absent:  '{text[:70]}'")

    # Check Table A cell
    for t in doc2.tables:
        tbl = t._tbl
        rows = tbl.findall(f".//{{{W}}}tr")
        if not rows:
            continue
        first_texts = [e.text or "" for e in rows[0].findall(f".//{{{W}}}t")]
        if "Method" in str(first_texts):
            for row in rows:
                for tc in row.findall(f".//{{{W}}}tc"):
                    cell = "".join(e.text or "" for e in tc.findall(f".//{{{W}}}t"))
                    if "No violations identified" in cell:
                        print("  [OK] Table A exchangeability cell updated")
                    if "exchangeability holds by construction" in cell:
                        print("  [FAIL] Table A old exchangeability claim still present")
                        all_ok = False

    print(f"\n{'All checks passed.' if all_ok else 'SOME CHECKS FAILED — review output above.'}")


if __name__ == "__main__":
    main()
