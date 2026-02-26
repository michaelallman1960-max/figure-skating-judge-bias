"""
fix_third_review.py
Six fixes from third ChatGPT review:

1. Abstract: "Eleven events" → "Nine events"
2. Conclusion: "11 outcome-determinative events" → "nine outcome-determinative events"
3. Table 3 caption: "global CDF" → "residual-label (isuimpact_residual_v1)"
4. Methodological notes: add sentence that LOJO covers all 144 events incl. reduced panels
5. Section 4: add one sentence on event-local inference vs database-wide summaries
6. Appendix A: soften "miscalibration" framing — no known-null simulation exists
7. Section 6.2: add one sentence defending points-valued pooling across row types
"""

import shutil
from pathlib import Path
from docx import Document

DOCX_PATH = Path("judge_bias_isu_judging_system.docx")
BACKUP_PATH = Path("judge_bias_isu_judging_system.docx.bak_third_review")


def backup():
    shutil.copy2(DOCX_PATH, BACKUP_PATH)
    print(f"Backup → {BACKUP_PATH}")


def replace_in_para(para, old, new, label=""):
    full = para.text
    if old not in full:
        raise ValueError(f"[{label}] text not found:\n  '{old[:100]}'")
    new_text = full.replace(old, new)
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def find_para(doc, substring):
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    raise ValueError(f"Paragraph containing '{substring[:60]}' not found")


# ──────────────────────────────────────────────────────────────
# FIX 1: Abstract — "Eleven events" → "Nine events"
# ──────────────────────────────────────────────────────────────

def fix_abstract_eleven(doc):
    p = find_para(doc, "Eleven events satisfy a two-part outcome-determinative criterion")
    replace_in_para(
        p,
        "Eleven events satisfy a two-part outcome-determinative criterion",
        "Nine events satisfy a two-part outcome-determinative criterion",
        "fix1 abstract eleven→nine"
    )
    print("  Abstract: 'Eleven events' → 'Nine events'.")


# ──────────────────────────────────────────────────────────────
# FIX 2: Conclusion — "11 outcome-determinative events" → "nine"
# ──────────────────────────────────────────────────────────────

def fix_conclusion_eleven(doc):
    p = find_para(doc, "11 outcome-determinative events")
    replace_in_para(
        p,
        "and 11 outcome-determinative events",
        "and nine outcome-determinative events",
        "fix2 conclusion 11→nine"
    )
    print("  Conclusion: '11 outcome-determinative events' → 'nine'.")


# ──────────────────────────────────────────────────────────────
# FIX 3: Table 3 caption — "global CDF" → method tag
# ──────────────────────────────────────────────────────────────

def fix_table3_caption(doc):
    p = find_para(doc, "global CDF")
    replace_in_para(
        p,
        "global CDF",
        "residual-label (isuimpact_residual_v1)",
        "fix3 table3 caption global CDF"
    )
    print("  Table 3 caption: 'global CDF' → 'residual-label (isuimpact_residual_v1)'.")


# ──────────────────────────────────────────────────────────────
# FIX 4: Methodological notes — add LOJO 144-events sentence
# ──────────────────────────────────────────────────────────────

def fix_methodological_notes_lojo(doc):
    p = find_para(doc, "Events with fewer than nine judges are excluded")
    replace_in_para(
        p,
        "Events with fewer than nine judges are excluded (n=2 events). "
        "All 142 analyzed events have panels of exactly nine judges.",
        "Events with fewer than nine judges are excluded from permutation inference "
        "(n=2 events). All 142 analyzed events have panels of exactly nine judges. "
        "LOJO is computed for all 144 events, including the two events with "
        "reduced panels (seven judges each).",
        "fix4 methodological notes LOJO 144"
    )
    print("  Methodological notes: added LOJO covers 144 events sentence.")


# ──────────────────────────────────────────────────────────────
# FIX 5: Section 4 — add event-local vs database-wide sentence
# ──────────────────────────────────────────────────────────────

def fix_section4_event_local(doc):
    p = find_para(doc, "carries no implication of intent.")
    replace_in_para(
        p,
        "and carries no implication of intent.",
        "and carries no implication of intent. "
        "Throughout this paper, inference is event-local: p-values and q-values "
        "are computed within each event independently, and database-wide figures "
        "(Section\u00a09) are aggregates across these independent event-local results.",
        "fix5 section 4 event-local"
    )
    print("  Section 4: added event-local vs database-wide clarification.")


# ──────────────────────────────────────────────────────────────
# FIX 6: Appendix A — soften "miscalibration" framing
# ──────────────────────────────────────────────────────────────

def fix_appendix_a_soften(doc):
    p = find_para(doc, "v1 diagnostic only")
    # Find the exact sentence — check both \u00a0 and plain space variants for Section ref
    old_v1 = None
    candidates = [
        # with non-breaking space before 9.1
        ("Empirically (v1 diagnostic only), the quantile null yields a nominal "
         "rejection rate of 17.78% at p \u2264 0.05 versus the 5.0% null expectation "
         "\u2014 far above the expected level under a well-calibrated test. By contrast, "
         "the adopted v2 residual-label method yields 25.4% nominal p \u2264 0.05 across "
         "the database (Section\u00a09.1), consistent with a null/alternative mixture "
         "rather than miscalibration."),
        # with plain space before 9.1
        ("Empirically (v1 diagnostic only), the quantile null yields a nominal "
         "rejection rate of 17.78% at p \u2264 0.05 versus the 5.0% null expectation "
         "\u2014 far above the expected level under a well-calibrated test. By contrast, "
         "the adopted v2 residual-label method yields 25.4% nominal p \u2264 0.05 across "
         "the database (Section 9.1), consistent with a null/alternative mixture "
         "rather than miscalibration."),
    ]
    for c in candidates:
        if c in p.text:
            old_v1 = c
            break
    if old_v1 is None:
        # Print what we have for debugging
        idx = p.text.find("v1 diagnostic only")
        raise ValueError(
            f"Appendix A v1 sentence not found verbatim. "
            f"Found context: ...{repr(p.text[max(0,idx-10):idx+300])}..."
        )

    new_v1 = (
        "Empirically (v1 diagnostic only), the quantile null yields a nominal "
        "rejection rate of 17.78% at p \u2264 0.05, compared with 25.4% for the "
        "adopted v2 residual-label method on the same data (Section\u00a09.1). "
        "Without a known-null simulation we cannot rigorously attribute this "
        "difference to Type I error inflation; the lower v1 rejection rate is "
        "consistent with the misaligned null testing a different quantity than the "
        "target estimand. The v1 null does not represent a credible "
        "\u2018no directional bias\u2019 baseline for this estimand."
    )
    replace_in_para(p, old_v1, new_v1, "fix6 appendix A soften")
    print("  Appendix A: softened 'miscalibration' to 'not a credible baseline'.")


# ──────────────────────────────────────────────────────────────
# FIX 7: Section 6.2 — add pooling justification sentence
# ──────────────────────────────────────────────────────────────

def fix_sec62_pooling_justification(doc):
    p = find_para(doc, "stratification by row category is a robustness check deferred to future work.")
    replace_in_para(
        p,
        "Note: \u0394 values from all row types (GOE elements and PCS components) "
        "are pooled in the permutation; stratification by row category is a "
        "robustness check deferred to future work.",
        "Note: \u0394 values from all row types (GOE elements and PCS components) "
        "are pooled in the permutation; because each \u0394 is already expressed "
        "in points after ISU scaling (GOE factor applied; PCS components scaled "
        "to their contribution), pooling treats each row as a points-valued "
        "contribution on a common scale. Stratification by row category is a "
        "robustness check deferred to future work.",
        "fix7 sec62 pooling justification"
    )
    print("  Section 6.2: added points-valued pooling justification sentence.")


# ──────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────

def main():
    backup()
    doc = Document(str(DOCX_PATH))

    print("\n--- Fix 1: Abstract 'Eleven' → 'Nine' ---")
    fix_abstract_eleven(doc)

    print("\n--- Fix 2: Conclusion '11' → 'nine' ---")
    fix_conclusion_eleven(doc)

    print("\n--- Fix 3: Table 3 caption 'global CDF' → method tag ---")
    fix_table3_caption(doc)

    print("\n--- Fix 4: Methodological notes — LOJO 144 sentence ---")
    fix_methodological_notes_lojo(doc)

    print("\n--- Fix 5: Section 4 — event-local vs database-wide ---")
    fix_section4_event_local(doc)

    print("\n--- Fix 6: Appendix A — soften miscalibration framing ---")
    fix_appendix_a_soften(doc)

    print("\n--- Fix 7: Section 6.2 — pooling justification ---")
    fix_sec62_pooling_justification(doc)

    doc.save(str(DOCX_PATH))
    print(f"\nSaved → {DOCX_PATH}")

    # Verification
    doc2 = Document(str(DOCX_PATH))
    all_text = " ".join(p.text for p in doc2.paragraphs)

    checks_present = [
        "Nine events satisfy a two-part outcome-determinative criterion",
        "nine outcome-determinative events",
        "residual-label (isuimpact_residual_v1)",
        "LOJO is computed for all 144 events",
        "inference is event-local",
        "not represent a credible",
        "points-valued contribution on a common scale",
    ]
    checks_absent = [
        "Eleven events satisfy",
        "11 outcome-determinative",
        "global CDF",
        "miscalibration",
        "far above the expected level under a well-calibrated test",
    ]

    print("\nVerification:")
    all_ok = True
    for text in checks_present:
        found = text in all_text
        status = "OK" if found else "FAIL"
        if not found:
            all_ok = False
        print(f"  [{status}] present: '{text[:70]}'")
    for text in checks_absent:
        found = text in all_text
        status = "OK" if not found else "FAIL"
        if found:
            all_ok = False
        print(f"  [{status}] absent:  '{text[:70]}'")

    print(f"\n{'All checks passed.' if all_ok else 'SOME CHECKS FAILED — review output above.'}")


if __name__ == "__main__":
    main()
