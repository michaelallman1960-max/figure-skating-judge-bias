"""
fix_sixth_review.py
Sixth-review fixes (items 1–5):

  1. Add Emerson, Seltzer & Lin (2009) to reference list (after Lee 2008)
  2. Rewrite stale "style-adjusted null" sentence in §10 Discussion (para 58)
  3. Add event-local BH-FDR clarification sentence to §9.1 (para 40)
  4. Add "LOJO" column to Table 3 (docx table index 5)
  5. Add GOE+PCS decomposition sentence to §8 worked example (para 33)
"""

import shutil
from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DOCX = Path("judge_bias_isu_judging_system.docx")
BAK  = Path("judge_bias_isu_judging_system.docx.bak_sixth_review")
W    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# ── helpers ─────────────────────────────────────────────────────────────────

def replace_in_para(para, old, new):
    """Replace old→new in paragraph. Zeroes all runs, puts text in run 0.
    Only safe for paragraphs that are plain text (no bold/italic to preserve)."""
    full = para.text
    if old not in full:
        raise ValueError(f"Text not found: {repr(old[:80])}")
    new_text = full.replace(old, new, 1)
    for i, run in enumerate(para.runs):
        run.text = new_text if i == 0 else ""


def insert_para_after(after_para, text):
    """Insert a new plain-text paragraph immediately after after_para,
    copying its paragraph style (but not run formatting)."""
    # Clone the paragraph element for style, then clear runs
    new_p = deepcopy(after_para._element)
    for r in new_p.findall(qn("w:r")):
        new_p.remove(r)
    for ins in new_p.findall(qn("w:ins")):
        new_p.remove(ins)
    # Build a plain run
    r_elem = OxmlElement("w:r")
    t_elem = OxmlElement("w:t")
    t_elem.text = text
    t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_elem.append(t_elem)
    new_p.append(r_elem)
    after_para._element.addnext(new_p)
    return new_p


def add_table_column(table, header_text, cell_texts):
    """Add a column to the right side of table. cell_texts[i] = data for row i+1."""
    tbl = table._tbl
    # Direct-child rows only
    rows = [ch for ch in tbl if ch.tag == f"{{{W}}}tr"]

    # Extend tblGrid
    tblGrid = tbl.find(f"{{{W}}}tblGrid")
    if tblGrid is not None:
        last_gc = tblGrid.findall(f"{{{W}}}gridCol")[-1]
        tblGrid.append(deepcopy(last_gc))

    all_texts = [header_text] + list(cell_texts)
    for ri, row in enumerate(rows):
        cells = [ch for ch in row if ch.tag == f"{{{W}}}tc"]
        new_tc = deepcopy(cells[-1])
        # Blank out all text elements
        for t_elem in new_tc.findall(f".//{{{W}}}t"):
            t_elem.text = ""
        # Set our text in the first <w:t>
        t_elems = new_tc.findall(f".//{{{W}}}t")
        if t_elems:
            t_elems[0].text = all_texts[ri] if ri < len(all_texts) else ""
        row.append(new_tc)


# ── main ────────────────────────────────────────────────────────────────────

def main():
    shutil.copy(DOCX, BAK)
    print(f"Backed up to {BAK}")

    doc = Document(DOCX)

    # ── Fix 1: Add Emerson et al. (2009) reference after Lee (2008) ─────────
    EMERSON_REF = (
        "Emerson, J. D., Seltzer, M., and Lin, D. (2009). "
        "Assessing Judging Bias: An Example From the 2000 Olympic Games. "
        "The American Statistician, 63(2), 124\u2013131."
    )
    # Find the Lee (2008) paragraph
    lee_para = None
    for p in doc.paragraphs:
        if p.text.startswith("Lee, J. (2008)"):
            lee_para = p
            break
    if lee_para is None:
        raise ValueError("Lee (2008) paragraph not found")
    # Check it's not already there
    for p in doc.paragraphs:
        if "Emerson" in p.text and "2009" in p.text:
            print("  ⚠ Emerson (2009) already present — skipping Fix 1")
            break
    else:
        insert_para_after(lee_para, EMERSON_REF)
        print("  ✓ Fix 1: Emerson (2009) reference inserted after Lee (2008)")

    # ── Fix 2: Rewrite stale "style-adjusted" limitation in para 58 ─────────
    OLD_SECOND = (
        "Second, the style-adjusted null preserves judge marginal distributions "
        "within categories but does not preserve all dependence structures, "
        "e.g., common-mode shifts on particular competitors. The approach is "
        "therefore conservative for some forms of collusion and may still "
        "attribute certain structured behaviors to competitor-specific differential."
    )
    NEW_SECOND = (
        "Second, the residual-label null assumes that pooled delta values are "
        "exchangeable across entries for each judge\u2013pair test. This holds "
        "when median-of-others neutralization removes the shared quality signal. "
        "If entry-specific variation persists in the residuals \u2014 for example, "
        "due to sequencing effects or element difficulty variation not captured "
        "by base values \u2014 the null could be mildly anti-conservative. The "
        "approach may also understate collusion effects that manifest as "
        "correlated delta patterns across entries rather than single-judge "
        "directional differentials."
    )
    found = False
    for p in doc.paragraphs:
        if OLD_SECOND in p.text:
            replace_in_para(p, OLD_SECOND, NEW_SECOND)
            print("  ✓ Fix 2: stale 'style-adjusted' limitation rewritten")
            found = True
            break
    if not found:
        raise ValueError("Fix 2: OLD_SECOND text not found in any paragraph")

    # ── Fix 3: Add event-local BH-FDR clarification to §9.1 ─────────────────
    OLD_91 = "20,213 pairs (7.4% of all tests) remained significant."
    NEW_91 = (
        "20,213 pairs (7.4% of all tests) remained significant. "
        "Note that BH-FDR is applied independently within each event; "
        "the 20,213 figure aggregates per-event corrected counts and "
        "does not reflect a single globally corrected procedure."
    )
    found = False
    for p in doc.paragraphs:
        if OLD_91 in p.text:
            replace_in_para(p, OLD_91, NEW_91)
            print("  ✓ Fix 3: §9.1 event-local BH-FDR clarification added")
            found = True
            break
    if not found:
        raise ValueError("Fix 3: §9.1 text not found")

    # ── Fix 4: Add LOJO column to Table 3 (docx table index 5) ──────────────
    # LOJO winner_changes (0=No, 1=Yes) for the 9 Table 3 rows in order:
    # 1. OWG 2026 J1        → Yes
    # 2. OWG 2022 J4        → Yes
    # 3. 4CC 2025 Pair J9   → Yes
    # 4. 4CC 2025 Dance J7  → No
    # 5. Europeans 2024 J1  → No
    # 6. GP Final 23/24 J6  → Yes
    # 7. GP Final 22/23 J1  → No
    # 8. Europeans 2022 J2  → No
    # 9. Europeans 2022† J3 → Yes
    lojo_values = ["Yes", "Yes", "Yes", "No", "No", "Yes", "No", "No", "Yes"]
    if len(doc.tables) <= 5:
        raise ValueError("Fix 4: Table 5 not found")
    t5 = doc.tables[5]
    # Check header row — verify it's the right table
    rows = [ch for ch in t5._tbl if ch.tag == f"{{{W}}}tr"]
    first_cells = rows[0].findall(f".//{{{W}}}tc")
    header = ["".join(e.text or "" for e in tc.findall(f".//{{{W}}}t")) for tc in first_cells]
    if "Competition" not in header[0]:
        raise ValueError(f"Fix 4: unexpected table header: {header}")
    if "LOJO" in header:
        print("  ⚠ LOJO column already present — skipping Fix 4")
    else:
        add_table_column(t5, "LOJO", lojo_values)
        print("  ✓ Fix 4: LOJO column added to Table 3")

    # ── Fix 5: Add GOE+PCS decomposition sentence to §8 worked example ──────
    # Append after "No other nationality grouping ... J1\u2019s impact column."
    OLD_33_TAIL = "No other nationality grouping shows a comparable unidirectional pattern in J1\u2019s impact column."
    NEW_33_TAIL = (
        "No other nationality grouping shows a comparable unidirectional "
        "pattern in J1\u2019s impact column. "
        "Decomposing J1\u2019s net +1.19\u2009pt margin shift by scoring component: "
        "element (GOE) rows account for approximately +1.01\u2009pts (85%) and "
        "programme component scores (PCS) account for approximately +0.20\u2009pts (17%) "
        "(FRA contribution: +0.16 from GOE, +0.10 from PCS; "
        "USA contribution: \u22120.85 from GOE, \u22120.10 from PCS; "
        "minor rounding from GOE integer discreteness in the neutralization step)."
    )
    found = False
    for p in doc.paragraphs:
        if OLD_33_TAIL in p.text:
            replace_in_para(p, OLD_33_TAIL, NEW_33_TAIL)
            print("  \u2713 Fix 5: GOE+PCS decomposition sentence added to \u00a78")
            found = True
            break
    if not found:
        raise ValueError("Fix 5: para 33 tail text not found")

    doc.save(DOCX)
    print(f"\nSaved {DOCX}")

    # ── Verify ───────────────────────────────────────────────────────────────
    print("\n── Verification ──")
    doc2 = Document(DOCX)

    checks = [
        ("Emerson 2009 present",     lambda d: any("Emerson" in p.text and "2009" in p.text for p in d.paragraphs)),
        ("residual-label null (fix2)", lambda d: any("residual-label null assumes" in p.text for p in d.paragraphs)),
        ("style-adjusted GONE",      lambda d: not any("style-adjusted null preserves" in p.text for p in d.paragraphs)),
        ("event-local BH note (fix3)", lambda d: any("aggregates per-event corrected counts" in p.text for p in d.paragraphs)),
        ("LOJO column in table",     lambda d: "LOJO" in ["".join(e.text or "" for e in tc.findall(f".//{{{W}}}t"))
                                                for tc in [ch for ch in [ch for ch in d.tables[5]._tbl if ch.tag == f"{{{W}}}tr"][0] if ch.tag == f"{{{W}}}tc"]]),
        ("GOE decomp sentence (fix5)", lambda d: any("GOE integer discreteness" in p.text for p in d.paragraphs)),
    ]
    all_ok = True
    for label, check_fn in checks:
        ok = check_fn(doc2)
        symbol = "✓" if ok else "✗"
        print(f"  {symbol} {label}")
        if not ok:
            all_ok = False

    if all_ok:
        print("\nAll 5 fixes verified.")
    else:
        print("\nSome checks FAILED — review above.")


if __name__ == "__main__":
    main()
