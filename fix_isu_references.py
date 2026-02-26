"""
fix_isu_references.py
Update the two incomplete ISU references in the paper (paragraphs 63 and 64).

Reference 2 (para 63): Add URL and clean up title — ISU Judging System page
Reference 3 (para 64): Add year, edition date, and URL — Special Regulations PDF

URLs verified 2026-02-26 at https://www.isu.org/figure-skating-rules/
"""

import shutil
from pathlib import Path
from docx import Document

DOCX = Path("judge_bias_isu_judging_system.docx")
BAK  = Path("judge_bias_isu_judging_system.docx.bak_isu_refs")

# ── helpers ────────────────────────────────────────────────────────────────

def replace_in_para(para, old, new):
    """Replace old→new in paragraph, preserving run structure where possible."""
    full = para.text
    if old not in full:
        raise ValueError(f"Text not found in para: {repr(old)}")
    # Zero out all runs, put new text in first run
    new_text = full.replace(old, new, 1)
    for i, run in enumerate(para.runs):
        run.text = new_text if i == 0 else ""

# ── main ───────────────────────────────────────────────────────────────────

def main():
    shutil.copy(DOCX, BAK)
    print(f"Backed up to {BAK}")

    doc = Document(DOCX)

    # ── Fix 1: Reference 2 — ISU Judging System (Panel of Judges / trimmed mean)
    OLD_REF2 = (
        "International Skating Union (ISU). Figure Skating Rules: Panel of Judges "
        "and trimmed mean description. ISU website (accessed 2026-02-23)."
    )
    NEW_REF2 = (
        "International Skating Union (ISU). ISU Judging System: Panel of Judges "
        "and trimmed mean. https://www.isu.org/figure-skating-rules/ "
        "(accessed 2026-02-26)."
    )

    # ── Fix 2: Reference 3 — Special Regulations & Technical Rules
    OLD_REF3 = (
        "International Skating Union (ISU). Special Regulations and Technical Rules: "
        "Single & Pair Skating and Ice Dance (current edition)."
    )
    NEW_REF3 = (
        "International Skating Union (ISU). (2024). Special Regulations & Technical "
        "Rules: Single & Pair Skating and Ice Dance. International Skating Union "
        "(27 Nov 2024). https://www.isu.org/figure-skating-rules/"
    )

    fixes = [
        ("Ref 2 — Panel of Judges / trimmed mean", OLD_REF2, NEW_REF2),
        ("Ref 3 — Special Regulations",             OLD_REF3, NEW_REF3),
    ]

    for label, old, new in fixes:
        found = False
        for para in doc.paragraphs:
            if old in para.text:
                replace_in_para(para, old, new)
                print(f"  ✓ {label}")
                found = True
                break
        if not found:
            # Show closest paragraph for diagnosis
            for i, para in enumerate(doc.paragraphs):
                if "International Skating Union" in para.text and "Panel" in para.text:
                    print(f"  DIAG para {i}: {repr(para.text)}")
            raise ValueError(f"NOT FOUND: {label}\nLooking for: {repr(old)}")

    doc.save(DOCX)
    print(f"\nSaved {DOCX}")

    # ── Verify ──────────────────────────────────────────────────────────────
    print("\n── Verification ──")
    doc2 = Document(DOCX)
    for para in doc2.paragraphs:
        if "isu.org/figure-skating-rules" in para.text:
            print(f"  ✓ {para.text}")

if __name__ == "__main__":
    main()
