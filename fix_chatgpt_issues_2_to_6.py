"""
fix_chatgpt_issues_2_to_6.py
Fixes all confirmed issues from ChatGPT review:
  2. Table 3: remove 6 rows with negative BiasPoints (wrong directional criterion)
     and update Section 9.3 counts (11→9 events, 7.7%→6.3%, 8→7 competitions,
     p≤0.007→p≤0.004, row-7 reference in Valieva note)
  3. Appendix A: rewrite technical description of quantile-null flaw
  4. LOJO arithmetic: clarify 1,288 panel-size breakdown
  5. Section 6.2: add sentence about pooled permutations
  6. Section 9.5: add "conditional on exchangeability" qualifier
  +  Add page numbers to footer throughout
"""

import shutil
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCX_PATH = Path("judge_bias_isu_judging_system.docx")
BACKUP_PATH = Path("judge_bias_isu_judging_system.docx.bak_chatgpt_fixes")


def backup():
    shutil.copy2(DOCX_PATH, BACKUP_PATH)
    print(f"Backup → {BACKUP_PATH}")


def replace_in_para(para, old, new, label=""):
    full = para.text
    if old not in full:
        raise ValueError(f"[{label}] replacement text not found:\n  '{old[:80]}'")
    new_text = full.replace(old, new)
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def find_para(doc, substring):
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    raise ValueError(f"Paragraph containing '{substring[:60]}' not found")


# ──────────────────────────────────────────────────────────────
# FIX 2a: Remove 6 rows from Table 3 (Table 5 in docx)
# ──────────────────────────────────────────────────────────────

# Rows to remove: identified by judge + competition content
ROWS_TO_REMOVE = [
    "GP Final 2024/25",   # Pair SP J8 (−0.75) and J9 (−1.00) — two rows
    "J6",                 # 4CC 2025 Ice Dance FD J6 (−0.60)
    "Europeans 2024",     # Pair FS J2 (−1.54)
    "J1\nEuropeans 2022", # Europeans 2022 Men SP J1 (−1.02)
    "J5",                 # Europeans 2022 Women FS J5 (−0.83)
]

# More precise: identify by judge column AND bias sign
# We'll remove any row where the Bias (pts) cell starts with "−"
def fix_table3(doc):
    """Remove rows with negative bias values from Table 5 (paper's Table 3).
    Table 5 was built with raw XML (no tblGrid), so we access via lxml directly."""
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    # Find the outcome-determinative table: Table 5 in docx, header starts "Competition"
    outcome_tbl_elem = None
    for t in doc.tables:
        tbl = t._tbl
        rows = tbl.findall(f".//{{{W}}}tr")
        if not rows:
            continue
        # Check first row first cell for "Competition"
        first_row_texts = [e.text or "" for e in rows[0].findall(f".//{{{W}}}t")]
        if first_row_texts and first_row_texts[0].strip() == "Competition":
            outcome_tbl_elem = tbl
            break

    assert outcome_tbl_elem is not None, "Could not find outcome-determinative table"

    rows = outcome_tbl_elem.findall(f".//{{{W}}}tr")

    # Find bias column index from header row
    header_cells = rows[0].findall(f".//{{{W}}}tc")
    bias_col = None
    for i, cell in enumerate(header_cells):
        cell_text = "".join(t.text or "" for t in cell.findall(f".//{{{W}}}t"))
        if "Bias" in cell_text:
            bias_col = i
            break
    assert bias_col is not None, "Could not find Bias column"

    removed = 0
    rows_to_delete = []
    for row in rows[1:]:  # skip header
        cells = row.findall(f".//{{{W}}}tc")
        if bias_col < len(cells):
            bias_text = "".join(t.text or "" for t in cells[bias_col].findall(f".//{{{W}}}t"))
            bias_text = bias_text.strip()
            if bias_text.startswith("\u2212") or bias_text.startswith("-"):
                rows_to_delete.append(row)

    for row in rows_to_delete:
        outcome_tbl_elem.remove(row)
        removed += 1

    remaining_rows = outcome_tbl_elem.findall(f".//{{{W}}}tr")
    remaining = len(remaining_rows) - 1  # minus header
    print(f"  Removed {removed} rows from Table 3 (negative BiasPoints).")
    print(f"  Table 3 now has {remaining} data rows.")
    return remaining


# ──────────────────────────────────────────────────────────────
# FIX 2b: Update Section 9.3 text counts
# ──────────────────────────────────────────────────────────────

def fix_sec93_counts(doc):
    # Para [46]: "Eleven of 142 events (7.7%)" → "Nine of 142 events (6.3%)"
    p46 = find_para(doc, "9.3. Outcome-determinative events")
    replace_in_para(p46,
        "Eleven of 142 events (7.7%) satisfy this criterion (Table 3).",
        "Nine of 142 events (6.3%) satisfy this criterion (Table 3).",
        "9.3 para46")
    print("  Updated 9.3 para: 'Eleven' → 'Nine', 7.7% → 6.3%.")

    # Para [47]: eleven→nine, eight→seven competitions, p/q ceiling
    p47 = find_para(doc, "The eleven events span")
    replace_in_para(p47,
        "The eleven events span all four main disciplines and eight separate competitions. "
        "Margins range from 0.24 to 1.33 points and bias magnitudes from 0.55 to 1.95 points; "
        "all are statistically strong (p \u2264 0.007, q \u2264 0.049).",
        "The nine events span all four main disciplines and seven separate competitions. "
        "Margins range from 0.24 to 1.33 points and bias magnitudes from 0.55 to 1.95 points; "
        "all are statistically strong (p \u2264 0.004, q \u2264 0.050).",
        "9.3 para47")
    print("  Updated 9.3 para: 'eleven'→'nine', 8→7 competitions, p/q ceilings.")

    # Para [48] Valieva note: "row 7" → "the final row"
    p48 = find_para(doc, "Kamila Valieva")
    replace_in_para(p48,
        "European Championships 2022 Women's Free Skating (Table 3, row 7)",
        "European Championships 2022 Women's Free Skating (Table 3, final row)",
        "Valieva row ref")
    print("  Updated Valieva note: 'row 7' → 'final row'.")


# ──────────────────────────────────────────────────────────────
# FIX 3: Rewrite Appendix A technical explanation
# ──────────────────────────────────────────────────────────────

def fix_appendix_a(doc):
    p79 = find_para(doc, "The quantile permutation null (isuimpact_quantile_v1) mapped")

    old_text = (
        "The quantile permutation null (isuimpact_quantile_v1) mapped each judge\u2019s marks "
        "to percentiles using that judge\u2019s empirical CDF, permuted percentile labels within "
        "each scoring row across judges, then mapped back via each judge\u2019s inverse CDF. "
        "The flaw: after permuting, judge j receives a percentile originally assigned to judge k "
        "and maps it through judge j\u2019s own CDF. This is valid only if both judges have "
        "identical style distributions \u2014 which they do not. The resulting null distribution "
        "is unrealistically narrow, inflating Type I errors approximately 355\u00d7 across the "
        "full dataset."
    )

    new_text = (
        "The quantile permutation null (isuimpact_quantile_v1) mapped each judge\u2019s marks "
        "to percentiles within that judge\u2019s empirical CDF, permuted percentile labels across "
        "judges within each scoring row, then mapped back via each judge\u2019s inverse CDF. "
        "The flaw is a misalignment of the permutation axis. The quantile null permutes "
        "across judges within a row \u2014 asking whether judge j\u2019s rank among judges "
        "for this element is unusual \u2014 but the target null is within judge, across "
        "competitors: whether judge j\u2019s deltas for competitor A versus competitor B are "
        "exchangeable. Although mapping through each judge\u2019s own inverse CDF does "
        "preserve judge j\u2019s marginal distribution, the exchangeability violation is "
        "structural: the null tests the wrong comparison object. Empirically, the quantile "
        "null yields nominal p \u2264 0.05 for 17.78% of pairwise tests versus the 5.0% "
        "expected under a well-calibrated null (3.6\u00d7 inflation), indicating systematic "
        "rejection of true-null cases."
    )

    if old_text not in p79.text:
        raise ValueError("Appendix A old text not found verbatim")

    full = p79.text.replace(old_text, new_text)
    for run in p79.runs:
        run.text = ""
    if p79.runs:
        p79.runs[0].text = full
    else:
        p79.add_run(full)
    print("  Rewrote Appendix A: corrected exchangeability flaw description, replaced 355×.")


# ──────────────────────────────────────────────────────────────
# FIX 4: LOJO arithmetic clarification
# ──────────────────────────────────────────────────────────────

def fix_lojo_arithmetic(doc):
    p52 = find_para(doc, "Across 1,288 judge-event computations")
    replace_in_para(p52,
        "Across 1,288 judge-event computations (judge-event pairs across all 144 events),",
        "Across 1,288 judge-event computations across all 144 events "
        "(138 events \u00d7 9 judges + 4 events \u00d7 8 judges + 2 events \u00d7 7 judges; "
        "the ISU Four Continents 2022 used reduced panels),",
        "LOJO arithmetic")
    print("  Updated LOJO arithmetic: clarified 1,288 panel-size breakdown.")


# ──────────────────────────────────────────────────────────────
# FIX 5: Section 6.2 — add pooled-permutations note
# ──────────────────────────────────────────────────────────────

def fix_sec62_pooled(doc):
    p22 = find_para(doc, "6.2. Residual-label (delta-exchange) permutation null")
    old_end = ("This follows Emerson, Seltzer & Lin (2009) and corrects the exchangeability "
               "violation in a prior quantile-based null (see Appendix A).")
    new_end = ("This follows Emerson, Seltzer & Lin (2009) and corrects the exchangeability "
               "violation in a prior quantile-based null (see Appendix A). "
               "Note: \u0394 values from all row types (GOE elements and PCS components) "
               "are pooled in the permutation; stratification by row category is a "
               "robustness check deferred to future work.")
    replace_in_para(p22, old_end, new_end, "6.2 pooled note")
    print("  Added pooled-permutations note to Section 6.2.")


# ──────────────────────────────────────────────────────────────
# FIX 6: Section 9.5 — calibration conditionality
# ──────────────────────────────────────────────────────────────

def fix_calibration_conditionality(doc):
    p53 = find_para(doc, "9.5. Stability, calibration, and power")
    replace_in_para(p53,
        "Calibration: residual-label permutation p-values are exact by construction",
        "Calibration: residual-label permutation p-values are exact conditional on "
        "the exchangeability assumption holding",
        "9.5 calibration conditionality")
    print("  Updated Section 9.5: added 'conditional on exchangeability' qualifier.")


# ──────────────────────────────────────────────────────────────
# ADD PAGE NUMBERS
# ──────────────────────────────────────────────────────────────

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        # Use existing paragraph or create one
        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()  # clear any existing content

        run = p.add_run()

        def fld(fld_type):
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), fld_type)
            return fc

        def instr(text):
            it = OxmlElement("w:instrText")
            it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            it.text = text
            return it

        run._r.append(fld("begin"))
        run._r.append(instr(" PAGE "))
        run._r.append(fld("end"))

    print("  Added page numbers to footer (all sections).")


# ──────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────

def main():
    backup()
    doc = Document(str(DOCX_PATH))

    print("\n--- Fix 2a: Remove negative-BiasPoints rows from Table 3 ---")
    n_remaining = fix_table3(doc)

    print("\n--- Fix 2b: Update Section 9.3 counts ---")
    fix_sec93_counts(doc)

    print("\n--- Fix 3: Rewrite Appendix A ---")
    fix_appendix_a(doc)

    print("\n--- Fix 4: LOJO arithmetic ---")
    fix_lojo_arithmetic(doc)

    print("\n--- Fix 5: Section 6.2 pooled-permutations note ---")
    fix_sec62_pooled(doc)

    print("\n--- Fix 6: Section 9.5 calibration conditionality ---")
    fix_calibration_conditionality(doc)

    print("\n--- Adding page numbers ---")
    add_page_numbers(doc)

    doc.save(str(DOCX_PATH))
    print(f"\nSaved → {DOCX_PATH}")

    # Verification
    doc2 = Document(str(DOCX_PATH))
    checks = [
        ("Nine of 142 events (6.3%)", True),
        ("nine events span all four main disciplines and seven", True),
        ("p \u2264 0.004, q \u2264 0.050", True),
        ("final row", True),
        ("misalignment of the permutation axis", True),
        ("3.6\u00d7 inflation", True),
        ("138 events \u00d7 9 judges", True),
        ("pooled in the permutation", True),
        ("exact conditional on the exchangeability", True),
        ("Eleven of 142", False),   # should NOT exist
        ("355\u00d7", False),        # should NOT exist
    ]
    print("\nVerification:")
    for text, should_exist in checks:
        found = any(text in p.text for p in doc2.paragraphs)
        ok = found == should_exist
        status = "OK" if ok else "FAIL"
        marker = "" if should_exist else "(absent)"
        print(f"  [{status}] '{text[:60]}' {marker}")

    # Table 3 row count
    for t in doc2.tables:
        try:
            if t.cell(0, 0).text.strip() == "Competition":
                data_rows = len(t.rows) - 1
                status = "OK" if data_rows == 9 else "FAIL"
                print(f"  [{status}] Table 3 has {data_rows} data rows (expected 9)")
                break
        except Exception:
            pass


if __name__ == "__main__":
    main()
