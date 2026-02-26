"""
fix_fourth_review.py
Six fixes from fourth ChatGPT review:

1. Section 8: Correct factual error — "not FDR-significant at 0.05" is wrong (J5 is q=0.035).
   Add J5 acknowledgment and correct the claim.
2. Section 6.2: Add dependence sentence — within-entry dependence affects power, not validity.
3. Section 7: Clarify margin(A,B) uses official published segment totals.
4. Section 9.5: Soften "approximately uniform density" → "upper tail broadly consistent with null."
5. Appendix A: Simplify — remove 17.78% comparison, keep conclusion + reference to diagnostics file.
6. Abstract: "null models that preserve judge style" → "retain judge-level scoring behavior."
"""

import shutil
from pathlib import Path
from docx import Document

DOCX_PATH = Path("judge_bias_isu_judging_system.docx")
BACKUP_PATH = Path("judge_bias_isu_judging_system.docx.bak_fourth_review")


def backup():
    shutil.copy2(DOCX_PATH, BACKUP_PATH)
    print(f"Backup → {BACKUP_PATH}")


def replace_in_para(para, old, new, label=""):
    full = para.text
    if old not in full:
        raise ValueError(f"[{label}] text not found:\n  '{old[:100]}'")
    new_text = full.replace(old, new)
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def find_para(doc, substring):
    for p in doc.paragraphs:
        if substring in p.text:
            return p
    raise ValueError(f"Paragraph containing '{substring[:60]}' not found")


# ──────────────────────────────────────────────────────────────
# FIX 1: Section 8 — correct J5 prose error and acknowledge
# "Other judges' shifts...not FDR-significant at 0.05" is FALSE:
# J5 has q=0.035 (significant), B_{J5}=-0.770 (toward silver).
# ──────────────────────────────────────────────────────────────

def fix_j5_acknowledgment(doc):
    p = find_para(doc, "Other judges\u2019 shifts for this pair are smaller in magnitude "
                        "and not FDR-significant at 0.05.")
    replace_in_para(
        p,
        "Other judges\u2019 shifts for this pair are smaller in magnitude "
        "and not FDR-significant at 0.05.",
        "Judge J5 (Janis Engel, GER) also yields a significant margin shift "
        "for this pair in the opposite direction (B_{J5}(A,B)=\u22120.770, q=0.035), "
        "reflecting a shift toward B (silver); this does not satisfy the directional "
        "outcome-determinative criterion and would widen the A\u2013B margin if J5 "
        "were removed. No other judge yields q \u2264 0.05 for this pair.",
        "fix1 J5 acknowledgment"
    )
    print("  Section 8: corrected J5 prose error; added J5 significance acknowledgment.")


# ──────────────────────────────────────────────────────────────
# FIX 2: Section 6.2 — add within-entry dependence sentence
# ──────────────────────────────────────────────────────────────

def fix_sec62_dependence(doc):
    p = find_para(doc, "Stratification by row category is a robustness check "
                        "deferred to future work.")
    replace_in_para(
        p,
        "Stratification by row category is a robustness check deferred to future work.",
        "Stratification by row category is a robustness check deferred to future work. "
        "While \u0394 values within an entry need not be independent (elements within "
        "a program share a competitor-quality context), the residual-label test requires "
        "only exchangeability of the pooled \u0394 multiset across entries under the "
        "null; within-entry dependence affects power and the discreteness of the "
        "permutation distribution, not test validity.",
        "fix2 sec62 dependence"
    )
    print("  Section 6.2: added within-entry dependence sentence.")


# ──────────────────────────────────────────────────────────────
# FIX 3: Section 7 — clarify margin uses official published totals
# ──────────────────────────────────────────────────────────────

def fix_margin_definition(doc):
    p = find_para(doc, "margin(A,B) is the observed gold")
    replace_in_para(
        p,
        "margin(A,B) is the observed gold\u2013silver score difference.",
        "margin(A,B) is the official published segment total difference between "
        "A and B (ISU-rounded; the bias B_j is computed on raw unrounded marks, "
        "so differences are at most 0.01 pts).",
        "fix3 margin definition"
    )
    print("  Section 7: clarified margin(A,B) uses official published segment totals.")


# ──────────────────────────────────────────────────────────────
# FIX 4: Section 9.5 — soften "approximately uniform density"
# ──────────────────────────────────────────────────────────────

def fix_sec95_uniform(doc):
    p = find_para(doc, "approximately uniform density above")
    replace_in_para(
        p,
        "a pronounced spike at small values and approximately uniform density "
        "above p\u2248\u200a0.05, the two-component signature of a mixture of "
        "true-null and true-alternative hypotheses.",
        "a pronounced spike at small values and an upper tail broadly consistent "
        "with the null above p\u2248\u200a0.05, the two-component signature of a "
        "mixture of true-null and true-alternative hypotheses.",
        "fix4 sec95 uniform soften"
    )
    print("  Section 9.5: 'approximately uniform density' → 'upper tail broadly consistent with null'.")


# ──────────────────────────────────────────────────────────────
# FIX 5: Appendix A — simplify empirical comparison, keep conclusion
# ──────────────────────────────────────────────────────────────

def fix_appendix_a_simplify(doc):
    p = find_para(doc, "Empirically (v1 diagnostic only)")
    # The block to replace (check both curly and straight quote variants)
    old_block = (
        "Empirically (v1 diagnostic only), the quantile null yields a nominal "
        "rejection rate of 17.78% at p \u2264 0.05, compared with 25.4% for the "
        "adopted v2 residual-label method on the same data (Section\u00a09.1). "
        "Without a known-null simulation we cannot rigorously attribute this "
        "difference to Type I error inflation; the lower v1 rejection rate is "
        "consistent with the misaligned null testing a different quantity than the "
        "target estimand. The v1 null does not represent a credible "
        "\u2018no directional bias\u2019 baseline for this estimand."
    )
    # Fallback with straight quotes if curly not found
    old_block_straight = (
        "Empirically (v1 diagnostic only), the quantile null yields a nominal "
        "rejection rate of 17.78% at p \u2264 0.05, compared with 25.4% for the "
        "adopted v2 residual-label method on the same data (Section\u00a09.1). "
        "Without a known-null simulation we cannot rigorously attribute this "
        "difference to Type I error inflation; the lower v1 rejection rate is "
        "consistent with the misaligned null testing a different quantity than the "
        "target estimand. The v1 null does not represent a credible "
        "'no directional bias' baseline for this estimand."
    )

    full = p.text
    if old_block in full:
        old = old_block
    elif old_block_straight in full:
        old = old_block_straight
    else:
        idx = full.find("Empirically (v1 diagnostic only)")
        raise ValueError(
            f"Appendix A block not found verbatim. Context: "
            f"...{repr(full[idx:idx+400])}..."
        )

    new_block = (
        "The v1 null therefore does not represent a credible "
        "\u2018no directional bias\u2019 baseline for this estimand. "
        "Empirical diagnostics (including nominal rejection rates under both "
        "methods) are reported in methodology_diagnosis_v1.md."
    )
    replace_in_para(p, old, new_block, "fix5 appendix A simplify")
    print("  Appendix A: removed 17.78% comparison; replaced with conclusion + diagnostics reference.")


# ──────────────────────────────────────────────────────────────
# FIX 6: Abstract — "preserve judge style" → "retain judge-level scoring behavior"
# ──────────────────────────────────────────────────────────────

def fix_abstract_null_phrasing(doc):
    # Fix Abstract occurrence (ends with period in the sentence)
    p_abstract = find_para(doc, "null models that preserve judge style.")
    replace_in_para(
        p_abstract,
        "null models that preserve judge style.",
        "null models that retain judge-level scoring behavior.",
        "fix6a abstract null phrasing"
    )
    print("  Abstract: 'preserve judge style' → 'retain judge-level scoring behavior'.")

    # Fix Introduction occurrence (followed by ", so that the test targets")
    p_intro = find_para(doc, "null models that preserve judge style, so that")
    replace_in_para(
        p_intro,
        "null models that preserve judge style, so that",
        "null models that retain judge-level scoring behavior, so that",
        "fix6b intro null phrasing"
    )
    print("  Introduction: 'preserve judge style' → 'retain judge-level scoring behavior'.")


# ──────────────────────────────────────────────────────────────
# MAIN
# ──────────────────────────────────────────────────────────────

def main():
    backup()
    doc = Document(str(DOCX_PATH))

    print("\n--- Fix 1: Section 8 — J5 acknowledgment ---")
    fix_j5_acknowledgment(doc)

    print("\n--- Fix 2: Section 6.2 — within-entry dependence sentence ---")
    fix_sec62_dependence(doc)

    print("\n--- Fix 3: Section 7 — margin definition ---")
    fix_margin_definition(doc)

    print("\n--- Fix 4: Section 9.5 — soften 'approximately uniform density' ---")
    fix_sec95_uniform(doc)

    print("\n--- Fix 5: Appendix A — simplify empirical block ---")
    fix_appendix_a_simplify(doc)

    print("\n--- Fix 6: Abstract — null models phrasing ---")
    fix_abstract_null_phrasing(doc)

    doc.save(str(DOCX_PATH))
    print(f"\nSaved → {DOCX_PATH}")

    # Verification
    doc2 = Document(str(DOCX_PATH))
    all_text = " ".join(p.text for p in doc2.paragraphs)

    checks_present = [
        "Judge J5 (Janis Engel, GER) also yields a significant margin shift",
        "No other judge yields q",
        "within-entry dependence affects power",
        "official published segment total difference",
        "upper tail broadly consistent with the null",
        "methodology_diagnosis_v1.md",
        "retain judge-level scoring behavior",
    ]
    checks_absent = [
        "not FDR-significant at 0.05",
        "approximately uniform density",
        "17.78% at p",
        "null models that preserve judge style",
        "gold\u2013silver score difference",
    ]

    print("\nVerification:")
    all_ok = True
    for text in checks_present:
        found = text in all_text
        status = "OK" if found else "FAIL"
        if not found:
            all_ok = False
        print(f"  [{status}] present: '{text[:70]}'")
    for text in checks_absent:
        found = text in all_text
        status = "OK" if not found else "FAIL"
        if found:
            all_ok = False
        print(f"  [{status}] absent:  '{text[:70]}'")

    print(f"\n{'All checks passed.' if all_ok else 'SOME CHECKS FAILED — review output above.'}")


if __name__ == "__main__":
    main()
